"""
Script membuat dokumen Word sub-bab 5.2.1.4 UML
DD Shoes POS System
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# Margin 4-3-3-3
section = doc.sections[0]
section.page_height = Cm(29.7)
section.page_width  = Cm(21.0)
section.left_margin   = Cm(4.0)
section.right_margin  = Cm(3.0)
section.top_margin    = Cm(3.0)
section.bottom_margin = Cm(3.0)

HEADER_FILL = 'BDD7EE'  # biru muda header tabel deskripsi
WHITE_FILL  = 'FFFFFF'

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for b in ('top','left','bottom','right','insideH','insideV'):
        border = OxmlElement(f'w:{b}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr.append(tblBorders)

def para(text='', bold=False, size=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
         sb=0, sa=6, first_indent=None):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(sb)
    pf.space_after  = Pt(sa)
    pf.line_spacing = Pt(24)
    if first_indent is not None:
        pf.first_line_indent = Cm(first_indent)
    if text:
        r = p.add_run(text)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(size)
        r.font.bold = bold
    return p

def heading(text, sb=12, sa=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    p.paragraph_format.line_spacing = Pt(24)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    r.font.bold = True
    return p

def caption(text, align=WD_ALIGN_PARAGRAPH.CENTER):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(8)
    p.paragraph_format.line_spacing = Pt(24)
    parts = text.split(' ', 2)
    # Bold: "Gambar 5.x" atau "Tabel 5.x", normal: sisanya
    if len(parts) >= 2:
        r1 = p.add_run(parts[0] + ' ' + parts[1] + ' ')
        r1.font.name = 'Times New Roman'
        r1.font.size = Pt(12)
        r1.font.bold = True
        if len(parts) == 3:
            r2 = p.add_run(parts[2])
            r2.font.name = 'Times New Roman'
            r2.font.size = Pt(12)
            r2.font.bold = False
    return p

def code_block(text):
    """Blok kode PlantUML dengan font monospace"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.line_spacing = Pt(16)
    p.paragraph_format.left_indent  = Cm(0.5)
    r = p.add_run(text)
    r.font.name = 'Courier New'
    r.font.size = Pt(8)
    return p

def note_plantuml(label):
    """Keterangan placeholder gambar"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.line_spacing = Pt(20)
    r = p.add_run(f'[{label}]')
    r.font.name = 'Times New Roman'
    r.font.size = Pt(10)
    r.font.italic = True
    r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    return p

def usecase_table(tabel_num, nama, aktor, pre, flow_lines, post):
    """Tabel deskripsi use case format 2 kolom"""
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_before = Pt(8)
    p_cap.paragraph_format.space_after  = Pt(4)
    p_cap.paragraph_format.line_spacing = Pt(24)
    r1 = p_cap.add_run(f'Tabel 5.{tabel_num} ')
    r1.font.name = 'Times New Roman'; r1.font.size = Pt(12); r1.font.bold = True
    r2 = p_cap.add_run(f'Deskripsi Use Case {nama}')
    r2.font.name = 'Times New Roman'; r2.font.size = Pt(12); r2.font.bold = False

    tbl = doc.add_table(rows=5, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(tbl)

    col_widths = [Cm(4.0), Cm(10.0)]
    rows_data = [
        ('Nama Use Case', nama),
        ('Aktor', aktor),
        ('Pre-kondisi', pre),
        ('Flow of Event', '\n'.join(flow_lines)),
        ('Post-kondisi', post),
    ]

    for i, (label, val) in enumerate(rows_data):
        row = tbl.rows[i]
        row.cells[0].width = col_widths[0]
        row.cells[1].width = col_widths[1]

        set_cell_bg(row.cells[0], HEADER_FILL)
        set_cell_bg(row.cells[1], WHITE_FILL)
        row.cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        row.cells[1].vertical_alignment = WD_ALIGN_VERTICAL.TOP

        for para_obj in row.cells[0].paragraphs: para_obj.clear()
        for para_obj in row.cells[1].paragraphs: para_obj.clear()

        p0 = row.cells[0].paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p0.paragraph_format.space_before = Pt(2)
        p0.paragraph_format.space_after  = Pt(2)
        r = p0.add_run(label)
        r.font.name = 'Times New Roman'; r.font.size = Pt(11); r.font.bold = True

        p1 = row.cells[1].paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p1.paragraph_format.space_before = Pt(2)
        p1.paragraph_format.space_after  = Pt(2)
        r1 = p1.add_run(val)
        r1.font.name = 'Times New Roman'; r1.font.size = Pt(11); r1.font.bold = False

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

# ============================================================
# 5.2.1.4 UML
# ============================================================
heading('5.2.1.4  Unified Modeling Language (UML)', sb=0)

# ============================================================
# a. USE CASE DIAGRAM
# ============================================================
heading('a.  Use Case Diagram')

para(
    'Use case diagram merupakan salah satu jenis diagram dalam Unified Modeling '
    'Language (UML) yang digunakan untuk menggambarkan interaksi antara pengguna '
    '(aktor) dengan sistem secara fungsional. Diagram ini mendeskripsikan apa yang '
    'dapat dilakukan oleh sistem dari sudut pandang pengguna, tanpa menjelaskan '
    'bagaimana sistem melakukannya secara teknis. Dalam sistem informasi Point of '
    'Sale (POS) DD Shoes Store, terdapat dua aktor utama yang berinteraksi dengan '
    'sistem, yaitu Admin dan Kasir.',
    first_indent=1.25
)

para(
    'Admin memiliki hak akses penuh terhadap seluruh fitur sistem meliputi '
    'pengelolaan data master, inventaris, laporan, dan manajemen pengguna. Kasir '
    'memiliki hak akses terbatas yang difokuskan pada aktivitas operasional harian '
    'seperti melakukan transaksi penjualan, melihat katalog produk, dan melakukan '
    'closing shift. Seluruh use case dalam sistem memiliki relasi <<include>> '
    'terhadap use case Login, yang berarti pengguna diwajibkan untuk melakukan '
    'autentikasi terlebih dahulu sebelum dapat mengakses fitur-fitur tersebut.',
    first_indent=1.25
)

para(
    'Adapun use case diagram sistem informasi POS DD Shoes Store dapat dilihat '
    'pada Gambar 5.3 berikut.',
    first_indent=1.25
)

note_plantuml('Gambar Use Case Diagram — render dari kode PlantUML/Mermaid')
caption('Gambar 5.3 Use Case Diagram DD Shoes Store')

para(
    'Berdasarkan use case diagram tersebut, sistem informasi POS DD Shoes Store '
    'memiliki 16 use case yang terbagi antara dua aktor. Admin dapat mengakses '
    'seluruh 16 use case, sedangkan Kasir hanya dapat mengakses 4 use case yang '
    'berkaitan dengan operasional kasir harian. Relasi <<include>> menunjukkan '
    'bahwa setiap use case memerlukan autentikasi login sebagai prasyarat, '
    'sementara relasi <<extend>> pada use case Cetak Laporan PDF menunjukkan '
    'bahwa fitur tersebut bersifat opsional dan hanya aktif ketika admin memilih '
    'untuk mencetak dari halaman Laporan Penjualan.',
    first_indent=1.25
)

para('Adapun deskripsi setiap use case dijelaskan pada tabel-tabel berikut.',
     first_indent=1.25)

# Tabel-tabel deskripsi use case
USE_CASES = [
    (7, 'Login', 'Admin, Kasir',
     'Pengguna telah memiliki akun dan membuka halaman login sistem.',
     ['1. Pengguna membuka halaman login (/login/).',
      '2. Sistem menampilkan form input username dan password.',
      '3. Pengguna mengisi username dan password kemudian mengklik tombol masuk.',
      '4. Sistem memvalidasi kredensial melalui mekanisme autentikasi Django.',
      '5. Jika valid, Admin diarahkan ke dashboard dan Kasir ke halaman POS.',
      '6. Jika tidak valid, sistem menampilkan pesan kesalahan.'],
     'Pengguna berhasil login dan dapat mengakses fitur sistem sesuai rolenya.'),

    (8, 'Lihat Dashboard', 'Admin',
     'Admin telah berhasil login ke dalam sistem.',
     ['1. Admin membuka halaman utama (/).',
      '2. Sistem menampilkan 4 kartu statistik: total stok aktif, produk terjual hari ini, pendapatan hari ini, dan jumlah barang masuk.',
      '3. Sistem menampilkan grafik tren pendapatan harian/mingguan/bulanan via AJAX.',
      '4. Sistem menampilkan 5 merek terlaris dan 5 produk dead stock.'],
     'Admin memperoleh gambaran ringkas kondisi operasional toko secara real-time.'),

    (9, 'Kelola Master Merek', 'Admin',
     'Admin telah berhasil login ke dalam sistem.',
     ['1. Admin membuka halaman merek (/brands/).',
      '2. Sistem menampilkan daftar merek beserta jumlah produk.',
      '3. Admin dapat menambah merek baru melalui modal form.',
      '4. Admin dapat mengubah nama merek yang sudah ada.',
      '5. Admin dapat menghapus merek yang tidak memiliki produk terkait.',
      '6. Sistem menyimpan perubahan dan menampilkan notifikasi sukses.'],
     'Data merek berhasil diperbarui dan tersedia sebagai referensi pada katalog produk.'),

    (10, 'Kelola Master Kategori', 'Admin',
     'Admin telah berhasil login ke dalam sistem.',
     ['1. Admin membuka halaman kategori (/categories/).',
      '2. Sistem menampilkan daftar kategori beserta jumlah produk aktif.',
      '3. Admin dapat menambah, mengubah, atau menghapus data kategori melalui modal form AJAX.',
      '4. Sistem menyimpan perubahan dan menampilkan notifikasi sukses.'],
     'Data kategori berhasil diperbarui dan tersedia sebagai referensi pada katalog produk.'),

    (11, 'Kelola Master Pemasok', 'Admin',
     'Admin telah berhasil login ke dalam sistem.',
     ['1. Admin membuka halaman pemasok (/suppliers/).',
      '2. Sistem menampilkan daftar pemasok beserta nomor telepon dan catatan.',
      '3. Admin dapat menambah, mengubah, atau menghapus data pemasok.',
      '4. Sistem menyimpan perubahan dan menampilkan notifikasi sukses.'],
     'Data pemasok berhasil diperbarui dan tersedia sebagai referensi pada pencatatan barang masuk.'),

    (12, 'Kelola Produk', 'Admin',
     'Admin telah berhasil login dan data merek serta kategori telah tersedia.',
     ['1. Admin membuka halaman katalog produk (/products/).',
      '2. Sistem menampilkan daftar produk dengan filter merek, kategori, dan status.',
      '3. Admin dapat mengubah data produk melalui modal form.',
      '4. Admin dapat mengaktifkan atau menonaktifkan status produk.',
      '5. Kode produk di-generate otomatis dengan format [Inisial Merek]-[Ukuran]-[Nomor Urut].',
      '6. Stok produk dikelola otomatis oleh sistem melalui Django Signals.'],
     'Data produk berhasil diperbarui dan tersedia di katalog sistem.'),

    (13, 'Kelola Barang Masuk', 'Admin',
     'Admin telah berhasil login dan data pemasok telah tersedia.',
     ['1. Admin membuka halaman barang masuk (/stockin/).',
      '2. Admin mengklik tombol catat barang masuk dan mengisi form header.',
      '3. Admin menambahkan baris item produk beserta jumlah dan harga beli.',
      '4. Sistem memeriksa apakah produk sudah ada berdasarkan nama, ukuran, dan kondisi.',
      '5. Jika produk belum ada, sistem membuat produk baru secara otomatis beserta kode produk.',
      '6. Setelah disimpan, sistem otomatis menambah stok produk melalui Django Signals.'],
     'Data penerimaan barang tersimpan dan stok produk terkait bertambah secara otomatis.'),

    (14, 'Penyesuaian Stok', 'Admin',
     'Admin telah berhasil login dan terdapat produk dengan stok lebih dari nol.',
     ['1. Admin membuka halaman penyesuaian stok (/adjustment/).',
      '2. Admin memilih produk yang akan disesuaikan.',
      '3. Admin mengisi jumlah pengurangan, alasan, dan catatan.',
      '4. Sistem mengurangi stok produk sesuai jumlah yang diinput.',
      '5. Jika stok mencapai nol, sistem otomatis menonaktifkan status produk.',
      '6. Sistem menyimpan catatan penyesuaian sebagai riwayat audit.'],
     'Stok produk berhasil disesuaikan dan catatan penyesuaian tersimpan dalam sistem.'),

    (15, 'Manajemen User', 'Admin',
     'Admin telah berhasil login ke dalam sistem.',
     ['1. Admin membuka halaman manajemen user (/users/).',
      '2. Sistem menampilkan daftar akun pengguna beserta role dan status.',
      '3. Admin dapat menambah akun kasir baru dengan mengisi username, password, role, dan telepon.',
      '4. Admin dapat mengubah data akun yang sudah ada.',
      '5. Admin dapat mengaktifkan atau menonaktifkan akun kasir.'],
     'Data akun pengguna berhasil diperbarui sesuai kebutuhan operasional toko.'),

    (16, 'Laporan Penjualan', 'Admin',
     'Admin telah berhasil login dan terdapat data transaksi dalam sistem.',
     ['1. Admin membuka halaman laporan penjualan (/sales-report/).',
      '2. Admin dapat memfilter berdasarkan rentang tanggal, kasir, metode pembayaran, dan merek.',
      '3. Sistem menampilkan daftar transaksi beserta total pendapatan dan total diskon.',
      '4. Admin dapat mengklik detail transaksi untuk melihat struk format thermal receipt.',
      '5. Admin dapat membatalkan transaksi (void) dengan menyertakan alasan.'],
     'Admin memperoleh informasi rekap penjualan sesuai filter yang diterapkan.'),

    (17, 'Cetak Laporan PDF', 'Admin',
     'Admin berada di halaman laporan penjualan dan telah menerapkan filter yang diinginkan.',
     ['1. Admin mengklik tombol Cetak/Ekspor PDF.',
      '2. Sistem mengambil data transaksi yang sedang ditampilkan.',
      '3. Sistem membuat dokumen laporan dalam format HTML print-ready.',
      '4. Browser membuka jendela cetak dengan tampilan laporan formal.',
      '5. Admin mencetak atau menyimpan sebagai PDF.'],
     'Laporan penjualan berhasil dicetak atau disimpan dalam format PDF.'),

    (18, 'Laporan Inventory', 'Admin',
     'Admin telah berhasil login dan terdapat data produk dengan stok lebih dari nol.',
     ['1. Admin membuka halaman laporan inventory (/inventory-report/).',
      '2. Admin dapat memfilter berdasarkan merek dan kategori.',
      '3. Sistem menampilkan daftar produk beserta stok saat ini.',
      '4. Sistem menghitung total nilai modal (Sigma harga beli x stok) dan total potensi pendapatan (Sigma harga jual x stok).'],
     'Admin memperoleh gambaran valuasi aset inventaris toko secara akurat.'),

    (19, 'Pantau Closing Kasir', 'Admin',
     'Admin telah berhasil login dan terdapat data closing dari kasir.',
     ['1. Admin membuka halaman closing admin (/closing-admin/).',
      '2. Admin dapat memfilter berdasarkan tanggal dan nama kasir.',
      '3. Sistem menampilkan rekap closing beserta selisih kas.',
      '4. Admin dapat membuka kembali data closing dengan mengisi alasan jika terdapat kesalahan.',
      '5. Sistem mencatat tindakan pembukaan kembali sebagai audit trail.'],
     'Admin berhasil memantau dan mengelola data closing seluruh kasir.'),

    (20, 'Melakukan Transaksi POS', 'Admin, Kasir',
     'Pengguna telah berhasil login dan terdapat produk aktif dengan stok tersedia.',
     ['1. Pengguna membuka halaman POS (/pos/).',
      '2. Sistem memuat data produk aktif secara langsung (pre-embedded JSON).',
      '3. Pengguna mencari produk berdasarkan nama, kode produk, merek, atau kategori.',
      '4. Pengguna mengklik kartu produk untuk menambahkan ke keranjang.',
      '5. Pengguna mengatur jumlah item dan menambahkan diskon (opsional).',
      '6. Pengguna memilih metode pembayaran (Tunai, Transfer BRI, atau QRIS).',
      '7. Pengguna mengklik tombol Proses Bayar.',
      '8. Sistem menyimpan transaksi, mengurangi stok via Django Signals, dan menampilkan struk.'],
     'Transaksi berhasil tersimpan, stok produk berkurang, dan struk digital ditampilkan.'),

    (21, 'Lihat Katalog Produk', 'Admin, Kasir',
     'Pengguna telah berhasil login ke dalam sistem.',
     ['1. Pengguna membuka halaman katalog produk.',
      '2. Pengguna dapat mencari produk berdasarkan nama, kode produk, atau ukuran.',
      '3. Pengguna dapat memfilter berdasarkan merek dan kategori.',
      '4. Sistem menampilkan daftar produk aktif beserta stok, harga, ukuran, dan kondisi secara read-only.'],
     'Pengguna memperoleh informasi ketersediaan produk untuk keperluan pelayanan pelanggan.'),

    (22, 'Closing Shift', 'Admin, Kasir',
     'Pengguna telah berhasil login dan telah melakukan aktivitas transaksi pada hari tersebut.',
     ['1. Kasir membuka halaman closing shift (/closing-kasir/).',
      '2. Sistem menampilkan ringkasan transaksi hari ini berdasarkan metode pembayaran.',
      '3. Kasir memasukkan jumlah kas aktual yang dihitung secara fisik.',
      '4. Sistem menghitung selisih antara kas fisik dan data sistem secara otomatis.',
      '5. Kasir menambahkan catatan (opsional) kemudian mengklik Submit dan Kunci.',
      '6. Sistem menyimpan data closing dengan status terkunci.'],
     'Data closing berhasil terkunci dan dapat dipantau oleh admin melalui halaman closing admin.'),

    (23, 'Logout', 'Admin, Kasir',
     'Pengguna telah berhasil login ke dalam sistem.',
     ['1. Pengguna mengklik menu logout pada navigasi.',
      '2. Sistem menghapus sesi pengguna.',
      '3. Sistem mengarahkan pengguna kembali ke halaman login.'],
     'Sesi pengguna berakhir dan akses ke sistem dicabut hingga login kembali.'),
]

for uc in USE_CASES:
    usecase_table(*uc)

# ============================================================
# b. ACTIVITY DIAGRAM
# ============================================================
heading('b.  Activity Diagram')

para(
    'Activity diagram merupakan diagram UML yang digunakan untuk menggambarkan '
    'alur kerja atau aktivitas dalam suatu proses bisnis secara berurutan. Diagram '
    'ini menampilkan setiap langkah yang dilakukan oleh aktor maupun sistem dalam '
    'menyelesaikan suatu use case, termasuk percabangan kondisi (decision) dan '
    'aktivitas paralel. Dalam sistem informasi POS DD Shoes Store, activity diagram '
    'disajikan dalam format swimlane yang membagi aktivitas berdasarkan pihak yang '
    'bertanggung jawab, yaitu pengguna (Admin/Kasir) dan Sistem POS.',
    first_indent=1.25
)

ACTIVITIES = [
    ('1)', 'Activity Diagram Login', '5.4',
     'Proses login merupakan gerbang utama yang harus dilalui oleh setiap pengguna '
     'sebelum dapat mengakses fitur sistem. Activity diagram login menggambarkan '
     'alur autentikasi pengguna mulai dari pembukaan halaman login hingga sistem '
     'memberikan respons berdasarkan validasi kredensial yang dimasukkan.',
     'Berdasarkan activity diagram login, proses autentikasi dimulai ketika '
     'pengguna membuka halaman login dan sistem merespons dengan menampilkan '
     'formulir input. Setelah pengguna mengisi username dan password kemudian '
     'mengklik tombol masuk, sistem melakukan validasi kredensial menggunakan '
     'mekanisme autentikasi bawaan Django. Apabila kredensial valid, sistem '
     'membuat sesi pengguna dan mengarahkan ke halaman yang sesuai berdasarkan '
     'role — Admin diarahkan ke dashboard dan Kasir diarahkan ke halaman POS. '
     'Apabila kredensial tidak valid, sistem menampilkan pesan kesalahan dan '
     'pengguna dapat mengulang proses pengisian.'),

    ('2)', 'Activity Diagram Transaksi POS', '5.5',
     'Proses transaksi penjualan merupakan aktivitas inti dalam sistem POS yang '
     'melibatkan interaksi intensif antara kasir dan sistem. Activity diagram ini '
     'menggambarkan seluruh alur mulai dari pembukaan halaman POS, pemilihan '
     'produk, pengaturan keranjang, proses pembayaran, hingga penerbitan struk '
     'digital.',
     'Berdasarkan activity diagram transaksi POS, proses dimulai ketika kasir '
     'membuka halaman POS dan sistem langsung memuat data produk aktif tanpa '
     'memerlukan permintaan AJAX tambahan. Kasir melakukan pencarian dan pemilihan '
     'produk, mengatur keranjang belanja, kemudian memproses pembayaran. Untuk '
     'metode tunai, sistem menghitung kembalian secara otomatis, sedangkan untuk '
     'Transfer BRI dan QRIS jumlah yang diterima ditetapkan sama dengan total '
     'pembayaran. Setelah transaksi diproses, sistem menyimpan data ke database '
     'dan secara otomatis mengurangi stok melalui mekanisme Django Signals, '
     'kemudian menampilkan struk thermal receipt.'),

    ('3)', 'Activity Diagram Barang Masuk', '5.6',
     'Proses pencatatan barang masuk merupakan mekanisme utama pengisian stok '
     'dalam sistem. Activity diagram ini menggambarkan alur kerja admin dalam '
     'mencatat penerimaan barang dari pemasok, termasuk proses pembuatan produk '
     'baru secara otomatis apabila produk belum terdaftar dalam sistem.',
     'Berdasarkan activity diagram barang masuk, admin mengisi formulir header '
     'penerimaan kemudian menambahkan baris item produk satu per satu. Sistem '
     'memeriksa setiap produk berdasarkan kombinasi nama, ukuran, dan kondisi. '
     'Apabila produk sudah terdaftar, sistem memperbarui harga apabila terdapat '
     'perubahan. Apabila produk belum terdaftar, sistem secara otomatis membuat '
     'entri produk baru sekaligus men-generate kode produk unik. Setelah data '
     'disimpan, Django Signals menjalankan proses penambahan stok secara otomatis.'),

    ('4)', 'Activity Diagram Closing Shift', '5.7',
     'Proses closing shift merupakan aktivitas akhir shift kasir yang bertujuan '
     'untuk melakukan rekonsiliasi antara data penjualan sistem dengan uang fisik '
     'yang ada di laci kasir. Activity diagram ini menggambarkan alur penutupan '
     'shift mulai dari pengecekan data sistem hingga penguncian laporan closing.',
     'Berdasarkan activity diagram closing shift, sistem secara otomatis menghitung '
     'dan menampilkan rekap transaksi hari tersebut berdasarkan metode pembayaran '
     'ketika kasir membuka halaman closing. Kasir menghitung uang fisik secara '
     'manual dan memasukkan jumlahnya ke dalam sistem. Sistem kemudian menghitung '
     'selisih secara otomatis dan menampilkan status rekonsiliasi. Setelah kasir '
     'mengklik tombol submit, data closing tersimpan dengan status terkunci '
     'sehingga tidak dapat diubah kembali tanpa intervensi admin.'),

    ('5)', 'Activity Diagram Penyesuaian Stok', '5.8',
     'Proses penyesuaian stok dilakukan ketika terdapat pengurangan stok yang '
     'tidak berasal dari jalur penjualan normal, seperti barang rusak, hilang, '
     'atau retur ke pemasok. Activity diagram ini menggambarkan alur kerja admin '
     'dalam mencatat dan memproses penyesuaian stok produk.',
     'Berdasarkan activity diagram penyesuaian stok, admin memilih produk, mengisi '
     'jumlah pengurangan beserta alasan, kemudian menyimpan data. Sistem melakukan '
     'validasi untuk memastikan jumlah pengurangan tidak melebihi stok yang '
     'tersedia. Apabila valid, sistem menyimpan catatan penyesuaian dengan nilai '
     'negatif dan mengurangi stok produk. Apabila stok produk mencapai nol setelah '
     'pengurangan, sistem secara otomatis mengubah status produk menjadi tidak '
     'aktif. Seluruh riwayat penyesuaian tersimpan sebagai catatan audit.'),

    ('6)', 'Activity Diagram Kelola Data Master', '5.9',
     'Pengelolaan data master merupakan aktivitas administratif yang dilakukan '
     'oleh Admin untuk memelihara data referensi yang digunakan oleh seluruh modul '
     'sistem. Data master yang dimaksud meliputi data merek sepatu, kategori '
     'produk, pemasok barang, dan akun pengguna. Keempat modul tersebut memiliki '
     'pola alur kerja yang identik mengikuti pola CRUD sehingga direpresentasikan '
     'dalam satu activity diagram generik.',
     'Berdasarkan activity diagram kelola data master, seluruh modul pengelolaan '
     'data master mengikuti pola yang seragam. Admin membuka halaman yang dituju '
     'kemudian memilih aksi yang akan dilakukan. Untuk penambahan data, sistem '
     'menampilkan modal form dan menyimpan data setelah validasi berhasil. Untuk '
     'pengubahan data, sistem memuat data yang sudah ada ke dalam form. Untuk '
     'penghapusan atau penonaktifan, sistem meminta konfirmasi terlebih dahulu. '
     'Setelah setiap aksi berhasil, sistem memperbarui tampilan daftar data secara '
     'otomatis tanpa perlu memuat ulang halaman.'),

    ('7)', 'Activity Diagram Katalog Produk', '5.10',
     'Halaman katalog produk dapat diakses oleh Admin dan Kasir dengan tingkat '
     'hak akses yang berbeda. Admin dapat melihat dan mengelola produk secara '
     'penuh, sedangkan Kasir hanya dapat melihat ketersediaan produk secara '
     'read-only untuk keperluan pelayanan pelanggan.',
     'Berdasarkan activity diagram katalog produk, pengguna dapat menelusuri '
     'produk menggunakan tiga metode filter yaitu pencarian berdasarkan nama atau '
     'kode produk, filter berdasarkan merek, dan filter berdasarkan kategori. '
     'Sistem memperbarui tampilan secara dinamis sesuai filter yang diterapkan. '
     'Perbedaan hak akses diterapkan di akhir alur — Admin dapat mengklik tombol '
     'edit untuk mengelola data produk, sedangkan Kasir hanya dapat melihat '
     'informasi produk sebagai referensi pelayanan pelanggan.'),

    ('8)', 'Activity Diagram Laporan Penjualan dan Laporan Inventory', '5.11',
     'Modul pelaporan merupakan fasilitas yang disediakan khusus bagi Admin untuk '
     'memantau kinerja penjualan dan kondisi inventaris toko. Terdapat dua jenis '
     'laporan yaitu laporan penjualan yang menampilkan rekap transaksi dan laporan '
     'inventory yang menampilkan valuasi aset stok.',
     'Berdasarkan activity diagram laporan, terdapat dua jalur utama sesuai jenis '
     'laporan yang dipilih admin. Pada jalur laporan penjualan, admin dapat '
     'menerapkan berbagai filter untuk mempersempit data yang ditampilkan, kemudian '
     'memilih aksi lanjutan berupa melihat detail struk, membatalkan transaksi, '
     'atau mencetak laporan. Pada jalur laporan inventory, sistem secara otomatis '
     'menghitung valuasi aset berupa total nilai modal dan total potensi pendapatan '
     'berdasarkan data stok terkini.'),
]

for num, title, fig_num, intro, desc in ACTIVITIES:
    heading(f'{num}  {title}')
    para(intro, first_indent=1.25)
    para(f'Adapun {title.lower()} dapat dilihat pada Gambar {fig_num} berikut.',
         first_indent=1.25)
    note_plantuml(f'Gambar {fig_num} — {title} (render dari PlantUML)')
    caption(f'Gambar {fig_num} {title}')
    para(desc, first_indent=1.25)

# ============================================================
# c. SEQUENCE DIAGRAM
# ============================================================
heading('c.  Sequence Diagram')

para(
    'Sequence diagram merupakan diagram UML yang menggambarkan interaksi antar '
    'objek dalam suatu sistem secara berurutan berdasarkan waktu. Diagram ini '
    'menunjukkan pesan-pesan yang dikirimkan antar objek untuk menyelesaikan suatu '
    'skenario tertentu. Dalam sistem informasi POS DD Shoes Store, sequence diagram '
    'disajikan dengan empat partisipan utama yaitu Admin/Kasir sebagai aktor, '
    'Browser sebagai antarmuka klien, Sistem sebagai lapisan logika bisnis (Django '
    'views), dan Database sebagai lapisan penyimpanan data.',
    first_indent=1.25
)

SEQUENCES = [
    ('1)', 'Sequence Diagram Login', '5.12',
     'Sequence diagram login menggambarkan urutan pesan yang terjadi antara '
     'pengguna, browser, sistem, dan database dalam proses autentikasi pengguna '
     'ke dalam sistem.',
     'Berdasarkan sequence diagram login, proses dimulai ketika pengguna mengakses '
     'halaman login melalui browser dengan metode GET. Sistem merespons dengan '
     'menampilkan halaman login. Setelah pengguna mengisi kredensial dan mengklik '
     'tombol masuk, browser mengirimkan permintaan POST ke sistem. Sistem melakukan '
     'autentikasi dengan memeriksa kredensial ke database. Apabila valid, sistem '
     'membuat sesi pengguna dan memeriksa role untuk menentukan halaman tujuan '
     'redirect. Apabila tidak valid, sistem mengembalikan halaman login dengan '
     'pesan kesalahan.'),

    ('2)', 'Sequence Diagram Transaksi POS', '5.13',
     'Sequence diagram transaksi POS menggambarkan urutan interaksi yang terjadi '
     'dalam proses penjualan mulai dari pembukaan halaman POS hingga penerbitan '
     'struk pembayaran.',
     'Berdasarkan sequence diagram transaksi POS, sistem menggunakan mekanisme '
     'pre-embedded JSON untuk memuat data produk pada saat halaman pertama kali '
     'dibuka. Pencarian produk dapat dilakukan berdasarkan nama maupun kode produk. '
     'Pengelolaan keranjang belanja dilakukan sepenuhnya di sisi klien menggunakan '
     'JavaScript. Saat kasir mengklik proses bayar, browser mengirimkan seluruh '
     'data transaksi dalam satu permintaan POST. Sistem menggunakan mekanisme '
     'SELECT FOR UPDATE untuk mengunci baris produk guna mencegah race condition. '
     'Setelah transaksi tersimpan, Django Signals secara otomatis mengurangi stok.'),

    ('3)', 'Sequence Diagram Barang Masuk', '5.14',
     'Sequence diagram barang masuk menggambarkan urutan interaksi dalam proses '
     'pencatatan penerimaan barang dari pemasok, termasuk mekanisme pembuatan '
     'produk baru secara otomatis dan pembaruan stok.',
     'Berdasarkan sequence diagram barang masuk, sistem memproses setiap baris '
     'item produk secara iteratif. Untuk setiap item, sistem memeriksa keberadaan '
     'produk berdasarkan kombinasi nama, ukuran, dan kondisi. Apabila produk belum '
     'terdaftar, sistem secara otomatis membuat entri produk baru dan men-generate '
     'kode produk unik menggunakan fungsi generate_product_code(). Setelah setiap '
     'StockInDetails tersimpan, Django Signals menjalankan penambahan stok dan '
     'pembaruan harga beli menggunakan metode Last In secara otomatis.'),

    ('4)', 'Sequence Diagram Closing Shift', '5.15',
     'Sequence diagram closing shift menggambarkan urutan interaksi dalam proses '
     'penutupan shift kasir dan rekonsiliasi kas harian.',
     'Berdasarkan sequence diagram closing shift, sistem menghitung total transaksi '
     'harian secara otomatis pada saat halaman dibuka. Perhitungan dilakukan di '
     'sisi server untuk memastikan akurasi data. Ketika kasir mensubmit closing, '
     'sistem melakukan perhitungan ulang di sisi server sebagai validasi tambahan '
     'sebelum menyimpan data. Sistem juga menangani dua kondisi berbeda yaitu '
     'pembuatan record closing baru atau pembaruan record yang sebelumnya telah '
     'dibuka kembali oleh admin.'),

    ('5)', 'Sequence Diagram Kelola Data Master', '5.16',
     'Sequence diagram kelola data master menggambarkan urutan interaksi dalam '
     'proses pengelolaan data referensi sistem yang mencakup merek, kategori, '
     'pemasok, dan manajemen user.',
     'Berdasarkan sequence diagram kelola data master, seluruh operasi CRUD '
     'dilakukan melalui mekanisme AJAX sehingga halaman tidak perlu dimuat ulang '
     'secara penuh. Terdapat tiga skenario utama yaitu penambahan data baru, '
     'pengubahan data yang sudah ada, dan penghapusan atau penonaktifan data. '
     'Untuk operasi penambahan dan pengubahan, sistem menggunakan endpoint yang '
     'sama yaitu /save/ dengan pembeda berupa keberadaan parameter id. Setiap '
     'operasi yang berhasil mengembalikan respons JSON yang digunakan browser '
     'untuk memperbarui tampilan secara dinamis.'),

    ('6)', 'Sequence Diagram Laporan Penjualan', '5.17',
     'Sequence diagram laporan penjualan menggambarkan urutan interaksi dalam '
     'proses pengaksesan, pemfilteran, dan pengelolaan data laporan transaksi '
     'penjualan oleh admin.',
     'Berdasarkan sequence diagram laporan penjualan, terdapat tiga skenario '
     'utama yang dapat dilakukan admin. Pertama, mengakses laporan dengan '
     'menerapkan filter yang relevan — sistem menghitung total pendapatan dan '
     'diskon secara agregat di sisi server. Kedua, melihat detail struk transaksi '
     'yang ditampilkan dalam format thermal receipt menggunakan data yang diambil '
     'secara lengkap melalui prefetch related. Ketiga, membatalkan transaksi (void) '
     'yang memerlukan pengembalian stok produk secara iteratif sebelum mengubah '
     'status transaksi menjadi void.'),

    ('7)', 'Sequence Diagram Penyesuaian Stok', '5.18',
     'Sequence diagram penyesuaian stok menggambarkan urutan interaksi yang '
     'terjadi dalam proses koreksi stok produk yang dilakukan oleh admin akibat '
     'kondisi di luar jalur penjualan normal.',
     'Berdasarkan sequence diagram penyesuaian stok, sistem menampilkan hanya '
     'produk yang memiliki stok lebih dari nol pada formulir penyesuaian. Setelah '
     'admin mengisi dan mengirimkan formulir, sistem melakukan validasi di sisi '
     'server untuk memastikan jumlah pengurangan tidak melebihi stok yang tersedia. '
     'Apabila valid, sistem menyimpan catatan penyesuaian dengan nilai negatif '
     'pada field quantity, kemudian memperbarui stok produk secara langsung. '
     'Apabila stok produk mencapai nol, sistem secara otomatis mengubah status '
     'produk menjadi tidak aktif.'),

    ('8)', 'Sequence Diagram Katalog Produk', '5.19',
     'Sequence diagram katalog produk menggambarkan urutan interaksi yang terjadi '
     'saat pengguna menelusuri dan mencari produk melalui halaman katalog.',
     'Berdasarkan sequence diagram katalog produk, sistem memuat data produk '
     'beserta data referensi merek dan kategori dalam satu permintaan awal. Fitur '
     'pencarian mendukung dua parameter sekaligus yaitu nama produk dan kode '
     'produk melalui kondisi ILIKE pada query database. Filter berdasarkan merek '
     'dan kategori dikirimkan sebagai parameter GET pada URL sehingga kondisi '
     'filter dapat disimpan dan dibagikan melalui URL. Terdapat percabangan hak '
     'akses di bagian akhir alur — Admin dapat mengakses modal edit produk, '
     'sedangkan Kasir tidak memiliki akses ke fungsi tersebut.'),
]

for num, title, fig_num, intro, desc in SEQUENCES:
    heading(f'{num}  {title}')
    para(intro, first_indent=1.25)
    para(f'Adapun {title.lower()} dapat dilihat pada Gambar {fig_num} berikut.',
         first_indent=1.25)
    note_plantuml(f'Gambar {fig_num} — {title} (render dari PlantUML)')
    caption(f'Gambar {fig_num} {title}')
    para(desc, first_indent=1.25)

# ============================================================
# d. CLASS DIAGRAM
# ============================================================
heading('d.  Class Diagram')

para(
    'Class diagram merupakan diagram UML yang menggambarkan struktur statis dari '
    'suatu sistem dengan menampilkan kelas-kelas yang ada, atribut-atribut yang '
    'dimiliki setiap kelas, serta relasi antar kelas. Dalam konteks pengembangan '
    'berbasis framework Django, class diagram merepresentasikan model-model data '
    'yang didefinisikan dalam file models.py. Setiap kelas pada diagram ini '
    'berkorespondensi langsung dengan tabel pada basis data yang telah '
    'dinormalisasi hingga tahap 3NF.',
    first_indent=1.25
)

para(
    'Sistem informasi POS DD Shoes Store memiliki 12 kelas utama yang saling '
    'berelasi. Relasi antar kelas menggunakan notasi multiplisitas standar UML '
    'yaitu 1 untuk satu entitas dan 0..* untuk nol atau banyak entitas. Relasi '
    'asosiasi ditandai dengan garis penghubung beserta label yang mendeskripsikan '
    'hubungan antar kelas tersebut.',
    first_indent=1.25
)

para(
    'Adapun class diagram sistem informasi POS DD Shoes Store dapat dilihat pada '
    'Gambar 5.20 berikut.',
    first_indent=1.25
)

note_plantuml('Gambar 5.20 — Class Diagram (render dari PlantUML)')
caption('Gambar 5.20 Class Diagram Sistem Informasi POS DD Shoes Store')

para(
    'Berdasarkan class diagram tersebut, sistem informasi POS DD Shoes Store '
    'terdiri dari 12 kelas yang saling berelasi membentuk struktur data yang '
    'terorganisasi. Kelas User merupakan kelas sentral yang berelasi dengan hampir '
    'seluruh kelas lainnya karena setiap aktivitas dalam sistem selalu melibatkan '
    'pengguna yang terautentikasi. Kelas UserProfile memiliki relasi one-to-one '
    'dengan User untuk menyimpan informasi role dan nomor telepon tambahan.',
    first_indent=1.25
)

para(
    'Kelas Products merupakan kelas terpenting dalam domain bisnis sistem ini '
    'karena menjadi titik perpotongan antara alur pengadaan barang dan alur '
    'penjualan. Kelas ini berelasi dengan Categories dan Brands sebagai data '
    'referensi, serta berelasi dengan StockInDetails untuk pencatatan penerimaan '
    'dan TransactionDetails untuk pencatatan penjualan. Atribut product_code pada '
    'kelas Products di-generate secara otomatis oleh sistem dengan format '
    '[Inisial Merek]-[Ukuran]-[Nomor Urut] untuk memastikan keunikan setiap produk.',
    first_indent=1.25
)

para(
    'Kelas StockIns dan StockInDetails membentuk pola header-detail untuk mencatat '
    'penerimaan barang dari pemasok, sedangkan kelas Transactions dan '
    'TransactionDetails membentuk pola yang sama untuk mencatat transaksi penjualan. '
    'Kelas CashClosings bertanggung jawab atas rekonsiliasi kas harian kasir dengan '
    'dilengkapi mekanisme audit trail melalui atribut unlocked_by, unlocked_at, '
    'unlock_reason, dan unlock_count. Kelas StockAdjustments berperan sebagai '
    'pencatat koreksi stok di luar jalur normal dengan menyimpan nilai negatif '
    'pada atribut adj_quantity sebagai penanda pengurangan stok.',
    first_indent=1.25
)

# ============================================================
# SIMPAN FILE
# ============================================================
output_path = (r'c:\Users\IRVAN SUSANTO\pos_ddshoes2\pos_ddshoes\docs'
               r'\Bab_5_2_1_4_UML.docx')
doc.save(output_path)
print(f'File berhasil dibuat: {output_path}')
