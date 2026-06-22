"""
Script membuat dokumen Word Activity Diagram dalam format TABEL
Format: No | Swimlane | Aktivitas | Simbol UML | Decision? | Ke Langkah
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
section = doc.sections[0]
section.page_height = Cm(29.7)
section.page_width  = Cm(21.0)
section.left_margin   = Cm(4.0)
section.right_margin  = Cm(3.0)
section.top_margin    = Cm(3.0)
section.bottom_margin = Cm(3.0)

DARK_HEADER = '1F3864'
WHITE_FILL  = 'FFFFFF'
LIGHT_FILL  = 'F2F2F2'

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for b in ('top','left','bottom','right','insideH','insideV'):
        border = OxmlElement(f'w:{b}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr.append(tblBorders)

def cell_text(cell, text, bold=False, size=10, color_hex=None,
              align=WD_ALIGN_PARAGRAPH.CENTER, bg=None):
    if bg:
        set_cell_bg(cell, bg)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for p in cell.paragraphs: p.clear()
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(str(text))
    r.font.name = 'Times New Roman'
    r.font.size = Pt(size)
    r.font.bold = bold
    if color_hex:
        r.font.color.rgb = RGBColor(
            int(color_hex[0:2],16),
            int(color_hex[2:4],16),
            int(color_hex[4:6],16)
        )

def para(text, bold=False, size=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
         sb=0, sa=6, first_indent=None):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    p.paragraph_format.line_spacing = Pt(24)
    if first_indent: p.paragraph_format.first_line_indent = Cm(first_indent)
    if text:
        r = p.add_run(text)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(size)
        r.font.bold = bold
    return p

def heading(text, sb=12):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.line_spacing = Pt(24)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    r.font.bold = True
    return p

def tabel_caption(num, title):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.line_spacing = Pt(24)
    r1 = p.add_run(f'Tabel {num} ')
    r1.font.name = 'Times New Roman'; r1.font.size = Pt(12); r1.font.bold = True
    r2 = p.add_run(title)
    r2.font.name = 'Times New Roman'; r2.font.size = Pt(12); r2.font.bold = False
    return p

def build_activity_table(caption_num, caption_title, rows_data):
    """
    rows_data: list of (swimlane, aktivitas, simbol, decision, ke_langkah)
    """
    tabel_caption(caption_num, caption_title)
    headers = ['No', 'Swimlane', 'Aktivitas', 'Simbol UML', 'Decision?', 'Ke Langkah']
    col_w = [Cm(1.0), Cm(2.5), Cm(4.5), Cm(2.5), Cm(2.8), Cm(2.2)]

    tbl = doc.add_table(rows=1+len(rows_data), cols=6)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(tbl)

    # Header row
    hrow = tbl.rows[0]
    for i, (h, w) in enumerate(zip(headers, col_w)):
        hrow.cells[i].width = w
        cell_text(hrow.cells[i], h, bold=True, size=11,
                  color_hex='FFFFFF', bg=DARK_HEADER)

    # Data rows
    for ri, (lane, akt, simbol, dec, ke) in enumerate(rows_data):
        row = tbl.rows[ri+1]
        bg = WHITE_FILL if ri % 2 == 0 else LIGHT_FILL
        for i, w in enumerate(col_w):
            row.cells[i].width = w
        cell_text(row.cells[0], str(ri+1), size=10, bg=bg)
        cell_text(row.cells[1], lane,      size=10, bg=bg)
        cell_text(row.cells[2], akt,       size=10, bg=bg,
                  align=WD_ALIGN_PARAGRAPH.CENTER)
        cell_text(row.cells[3], simbol,    size=10, bg=bg)
        cell_text(row.cells[4], dec,       size=10, bg=bg)
        cell_text(row.cells[5], ke,        size=10, bg=bg)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

# ============================================================
# JUDUL SUB-BAB
# ============================================================
heading('b.  Activity Diagram', sb=0)

para(
    'Activity diagram merupakan diagram UML yang digunakan untuk menggambarkan '
    'alur kerja atau aktivitas dalam suatu proses bisnis secara berurutan. '
    'Dalam sistem informasi POS DD Shoes Store, activity diagram disajikan dalam '
    'format tabel swimlane yang mendeskripsikan setiap langkah aktivitas beserta '
    'simbol UML, percabangan keputusan, dan urutan transisi ke langkah berikutnya.',
    first_indent=1.25
)

# ============================================================
# 1. ACTIVITY LOGIN
# ============================================================
heading('1)  Activity Diagram Login')
para('Proses login untuk Admin dan Kasir.', first_indent=1.25)

build_activity_table(1, 'Langkah-langkah Activity Login', [
    ('Actor',  'Mulai',                               '● Start Node',  '-',                    '2'),
    ('Actor',  'Membuka halaman login',               '□ Action',      '-',                    '3'),
    ('Sistem', 'Menampilkan form login',              '□ Action',      '-',                    '4'),
    ('Actor',  'Mengisi username & password',         '□ Action',      '-',                    '5'),
    ('Actor',  'Mengklik tombol Masuk',               '□ Action',      '-',                    '6'),
    ('Sistem', 'Memvalidasi kredensial (Django authenticate)', '◇ Decision', 'Valid → 7 | Tidak Valid → 8', ''),
    ('Sistem', 'Cek role pengguna',                   '□ Action',      'Admin → 7a | Kasir → 7b', ''),
    ('Sistem', 'Redirect ke /dashboard/ (Admin)',     '□ Action',      '-',                    '9'),
    ('Sistem', 'Redirect ke /pos/ (Kasir)',           '□ Action',      '-',                    '9'),
    ('Sistem', 'Menampilkan pesan kesalahan',         '□ Action',      '-',                    '4'),
    ('Actor',  'Selesai',                             '◉ End Node',    '-',                    '-'),
])

para(
    'Berdasarkan tabel di atas, proses autentikasi dimulai dari langkah 1 '
    'sebagai titik awal, dilanjutkan dengan pembukaan halaman login oleh '
    'pengguna, validasi kredensial oleh sistem, dan diakhiri dengan redirect '
    'ke halaman yang sesuai berdasarkan role pengguna.',
    first_indent=1.25
)

# ============================================================
# 2. ACTIVITY TRANSAKSI POS
# ============================================================
heading('2)  Activity Diagram Transaksi POS')
para('Proses inti POS — berlaku untuk Admin dan Kasir.', first_indent=1.25)

build_activity_table(2, 'Langkah-langkah Activity Transaksi Penjualan', [
    ('Actor',  'Mulai',                                    '● Start Node', '-',                     '2'),
    ('Actor',  'Membuka halaman POS (/pos/)',              '□ Action',     '-',                     '3'),
    ('Sistem', 'Memuat produk aktif sebagai JSON',         '□ Action',     '-',                     '4'),
    ('Sistem', 'Menampilkan produk dalam kartu grid',      '□ Action',     '-',                     '5'),
    ('Actor',  'Mencari produk (nama/kode/merek)',         '□ Action',     '-',                     '6'),
    ('Actor',  'Mengklik kartu produk',                   '□ Action',     '-',                     '7'),
    ('Sistem', 'Menambah produk ke keranjang',            '□ Action',     '-',                     '8'),
    ('Actor',  'Mengatur jumlah (qty)',                   '□ Action',     '-',                     '9'),
    ('Actor',  'Ada item yang dihapus?',                  '◇ Decision',   'Ya → 10 | Tidak → 11',  ''),
    ('Actor',  'Menghapus item dari keranjang',           '□ Action',     '-',                     '8'),
    ('Actor',  'Memasukkan diskon (opsional)',             '□ Action',     '-',                     '12'),
    ('Sistem', 'Menghitung subtotal, diskon, grand total','□ Action',     '-',                     '13'),
    ('Actor',  'Mengklik tombol Proses Bayar',            '□ Action',     '-',                     '14'),
    ('Actor',  'Memilih metode pembayaran',               '◇ Decision',   'Tunai → 15 | Non-Tunai → 16', ''),
    ('Actor',  'Memasukkan jumlah uang diterima',         '□ Action',     '-',                     '17'),
    ('Sistem', 'Menetapkan diterima = grand total',       '□ Action',     '-',                     '17'),
    ('Sistem', 'Menyimpan Transactions & TransactionDetails', '□ Action', '-',                     '18'),
    ('Sistem', 'Mengurangi stok via Django Signals',      '□ Action',     '-',                     '19'),
    ('Sistem', 'Menampilkan struk thermal receipt',       '□ Action',     '-',                     '20'),
    ('Actor',  'Selesai',                                 '◉ End Node',   '-',                     '-'),
])

para(
    'Berdasarkan tabel di atas, alur transaksi POS mencakup 20 langkah yang '
    'melibatkan interaksi antara kasir dan sistem secara bergantian. Percabangan '
    'terjadi pada langkah 9 (penghapusan item) dan langkah 14 (metode pembayaran). '
    'Stok produk dikurangi secara otomatis oleh Django Signals pada langkah 18.',
    first_indent=1.25
)

# ============================================================
# 3. ACTIVITY BARANG MASUK
# ============================================================
heading('3)  Activity Diagram Barang Masuk')
para('Proses pencatatan penerimaan barang dari pemasok oleh Admin.', first_indent=1.25)

build_activity_table(3, 'Langkah-langkah Activity Barang Masuk', [
    ('Actor',  'Mulai',                                       '● Start Node', '-',                        '2'),
    ('Actor',  'Membuka halaman Barang Masuk (/stockin/)',    '□ Action',     '-',                        '3'),
    ('Sistem', 'Menampilkan riwayat penerimaan barang',       '□ Action',     '-',                        '4'),
    ('Actor',  'Mengklik tombol Catat Barang Masuk',          '□ Action',     '-',                        '5'),
    ('Sistem', 'Menampilkan form penerimaan barang',          '□ Action',     '-',                        '6'),
    ('Actor',  'Mengisi header (pemasok, tanggal, catatan)',  '□ Action',     '-',                        '7'),
    ('Actor',  'Menambahkan baris item produk',               '□ Action',     '-',                        '8'),
    ('Actor',  'Mengklik tombol Simpan',                      '□ Action',     '-',                        '9'),
    ('Sistem', 'Memeriksa produk (nama+ukuran+kondisi)',      '◇ Decision',   'Ada → 10 | Tidak Ada → 11', ''),
    ('Sistem', 'Memperbarui harga jika berubah',              '□ Action',     '-',                        '12'),
    ('Sistem', 'Membuat produk baru + generate kode produk',  '□ Action',     '-',                        '12'),
    ('Sistem', 'INSERT StockIns & StockInDetails',            '□ Action',     '-',                        '13'),
    ('Sistem', 'Tambah stok via Django Signals',              '□ Action',     '-',                        '14'),
    ('Sistem', 'Redirect ke daftar barang masuk + notifikasi','□ Action',     '-',                        '15'),
    ('Actor',  'Selesai',                                     '◉ End Node',   '-',                        '-'),
])

para(
    'Berdasarkan tabel di atas, percabangan utama terjadi pada langkah 9 '
    'yaitu pengecekan keberadaan produk berdasarkan kombinasi nama, ukuran, '
    'dan kondisi. Apabila produk belum ada, sistem secara otomatis membuat '
    'produk baru beserta kode produk unik pada langkah 11.',
    first_indent=1.25
)

# ============================================================
# 4. ACTIVITY CLOSING SHIFT
# ============================================================
heading('4)  Activity Diagram Closing Shift')
para('Proses penutupan shift dan rekonsiliasi kas oleh Kasir.', first_indent=1.25)

build_activity_table(4, 'Langkah-langkah Activity Closing Shift', [
    ('Actor',  'Mulai',                                         '● Start Node', '-',                    '2'),
    ('Actor',  'Membuka halaman Closing Shift (/closing-kasir/)','□ Action',    '-',                    '3'),
    ('Sistem', 'Menghitung total transaksi per metode bayar',    '□ Action',    '-',                    '4'),
    ('Sistem', 'Menampilkan ringkasan: Tunai, Transfer, QRIS',   '□ Action',    '-',                    '5'),
    ('Actor',  'Menghitung uang fisik di laci kasir',            '□ Action',    '-',                    '6'),
    ('Actor',  'Memasukkan jumlah kas fisik aktual',             '□ Action',    '-',                    '7'),
    ('Sistem', 'Menghitung selisih (kas fisik - total tunai)',    '◇ Decision',  'Selisih=0 → 8 | Selisih≠0 → 9', ''),
    ('Sistem', 'Menampilkan status: MATCH',                      '□ Action',    '-',                    '10'),
    ('Sistem', 'Menampilkan status: SELISIH + nilai perbedaan',  '□ Action',    '-',                    '10'),
    ('Actor',  'Menambahkan catatan (opsional)',                  '□ Action',    '-',                    '11'),
    ('Actor',  'Mengklik tombol Submit & Kunci',                 '□ Action',    '-',                    '12'),
    ('Sistem', 'Menyimpan CashClosings (is_locked=True)',        '□ Action',    '-',                    '13'),
    ('Sistem', 'Menampilkan notifikasi closing dikunci',         '□ Action',    '-',                    '14'),
    ('Actor',  'Selesai',                                        '◉ End Node',  '-',                    '-'),
])

para(
    'Berdasarkan tabel di atas, percabangan terjadi pada langkah 7 yaitu '
    'pengecekan selisih kas. Baik kondisi MATCH maupun SELISIH, kedua jalur '
    'bertemu kembali pada langkah 10 dan alur dilanjutkan hingga data '
    'closing berhasil dikunci oleh sistem.',
    first_indent=1.25
)

# ============================================================
# 5. ACTIVITY PENYESUAIAN STOK
# ============================================================
heading('5)  Activity Diagram Penyesuaian Stok')
para('Proses koreksi stok di luar jalur penjualan normal oleh Admin.', first_indent=1.25)

build_activity_table(5, 'Langkah-langkah Activity Penyesuaian Stok', [
    ('Actor',  'Mulai',                                          '● Start Node', '-',                        '2'),
    ('Actor',  'Membuka halaman Penyesuaian Stok (/adjustment/)','□ Action',     '-',                        '3'),
    ('Sistem', 'Menampilkan riwayat penyesuaian stok',           '□ Action',     '-',                        '4'),
    ('Actor',  'Mengklik tombol Tambah Penyesuaian',             '□ Action',     '-',                        '5'),
    ('Sistem', 'Menampilkan form penyesuaian stok',              '□ Action',     '-',                        '6'),
    ('Actor',  'Memilih produk yang akan disesuaikan',           '□ Action',     '-',                        '7'),
    ('Actor',  'Mengisi jumlah pengurangan',                     '□ Action',     '-',                        '8'),
    ('Actor',  'Memilih alasan (Rusak/Hilang/Retur/Lainnya)',    '□ Action',     '-',                        '9'),
    ('Actor',  'Mengklik tombol Simpan',                         '□ Action',     '-',                        '10'),
    ('Sistem', 'Memvalidasi jumlah vs stok saat ini',            '◇ Decision',   'Valid → 11 | Tidak Valid → 13', ''),
    ('Sistem', 'INSERT StockAdjustments (quantity negatif)',      '□ Action',     '-',                        '12'),
    ('Sistem', 'UPDATE Products: stock -= quantity',             '□ Action',     '-',                        '13'),
    ('Sistem', 'Stok = 0?',                                      '◇ Decision',   'Ya → 14 | Tidak → 15',    ''),
    ('Sistem', 'UPDATE Products: status = inactive',             '□ Action',     '-',                        '15'),
    ('Sistem', 'Menampilkan notifikasi berhasil',                '□ Action',     '-',                        '16'),
    ('Sistem', 'Menampilkan pesan error (jumlah melebihi stok)', '□ Action',     '-',                        '7'),
    ('Actor',  'Selesai',                                        '◉ End Node',   '-',                        '-'),
])

para(
    'Berdasarkan tabel di atas, terdapat dua percabangan utama yaitu pada '
    'langkah 10 (validasi jumlah) dan langkah 13 (pengecekan stok nol). '
    'Apabila validasi gagal, alur kembali ke langkah 7 untuk pengisian ulang. '
    'Apabila stok mencapai nol, sistem secara otomatis menonaktifkan produk.',
    first_indent=1.25
)

# ============================================================
# 6. ACTIVITY KELOLA DATA MASTER
# ============================================================
heading('6)  Activity Diagram Kelola Data Master')
para('Proses CRUD data master (Merek, Kategori, Pemasok, User) oleh Admin.', first_indent=1.25)

build_activity_table(6, 'Langkah-langkah Activity Kelola Data Master', [
    ('Actor',  'Mulai',                                        '● Start Node', '-',                              '2'),
    ('Actor',  'Membuka halaman data master',                  '□ Action',     '-',                              '3'),
    ('Sistem', 'Menampilkan daftar data beserta informasi',    '□ Action',     '-',                              '4'),
    ('Actor',  'Memilih aksi yang akan dilakukan',             '◇ Decision',   'Tambah → 5 | Ubah → 9 | Hapus → 13', ''),
    ('Actor',  'Mengklik tombol Tambah',                       '□ Action',     '-',                              '6'),
    ('Sistem', 'Menampilkan modal form penambahan data',       '□ Action',     '-',                              '7'),
    ('Actor',  'Mengisi data & mengklik Simpan',               '□ Action',     '-',                              '8'),
    ('Sistem', 'INSERT data baru ke database',                 '□ Action',     '-',                              '16'),
    ('Actor',  'Mengklik tombol Edit pada baris data',         '□ Action',     '-',                              '10'),
    ('Sistem', 'Menampilkan modal form dengan data existing',  '□ Action',     '-',                              '11'),
    ('Actor',  'Mengubah data & mengklik Simpan',              '□ Action',     '-',                              '12'),
    ('Sistem', 'UPDATE data di database',                      '□ Action',     '-',                              '16'),
    ('Actor',  'Mengklik tombol Hapus/Nonaktifkan',            '□ Action',     '-',                              '14'),
    ('Sistem', 'Menampilkan konfirmasi',                       '□ Action',     '-',                              '15'),
    ('Actor',  'Mengkonfirmasi aksi',                          '□ Action',     '-',                              '16'),
    ('Sistem', 'Memperbarui tampilan daftar data (AJAX)',      '□ Action',     '-',                              '17'),
    ('Sistem', 'Menampilkan notifikasi sukses',                '□ Action',     '-',                              '18'),
    ('Actor',  'Selesai',                                      '◉ End Node',   '-',                              '-'),
])

para(
    'Berdasarkan tabel di atas, percabangan utama terjadi pada langkah 4 '
    'yaitu pemilihan aksi CRUD. Ketiga jalur (Tambah, Ubah, Hapus) '
    'bertemu kembali pada langkah 16 yaitu pembaruan tampilan secara '
    'otomatis melalui mekanisme AJAX.',
    first_indent=1.25
)

# ============================================================
# 7. ACTIVITY KATALOG PRODUK
# ============================================================
heading('7)  Activity Diagram Katalog Produk')
para('Proses penelusuran dan pengelolaan katalog produk untuk Admin dan Kasir.', first_indent=1.25)

build_activity_table(7, 'Langkah-langkah Activity Katalog Produk', [
    ('Actor',  'Mulai',                                        '● Start Node', '-',                           '2'),
    ('Actor',  'Membuka halaman Katalog Produk',               '□ Action',     '-',                           '3'),
    ('Sistem', 'Menampilkan seluruh produk aktif beserta kode','□ Action',     '-',                           '4'),
    ('Actor',  'Memilih metode penelusuran',                   '◇ Decision',   'Cari → 5 | Filter Merek → 7 | Filter Kategori → 9', ''),
    ('Actor',  'Mengetik kata kunci pada kolom pencarian',     '□ Action',     '-',                           '6'),
    ('Sistem', 'Filter produk: nama ILIKE % atau kode ILIKE %','□ Action',     '-',                           '11'),
    ('Actor',  'Memilih merek pada dropdown filter',           '□ Action',     '-',                           '8'),
    ('Sistem', 'Filter produk berdasarkan brand_id',           '□ Action',     '-',                           '11'),
    ('Actor',  'Memilih kategori pada dropdown filter',        '□ Action',     '-',                           '10'),
    ('Sistem', 'Filter produk berdasarkan category_id',        '□ Action',     '-',                           '11'),
    ('Sistem', 'Memperbarui tampilan daftar/kartu produk',     '□ Action',     '-',                           '12'),
    ('Actor',  'Role pengguna?',                               '◇ Decision',   'Admin → 13 | Kasir → 15',     ''),
    ('Actor',  'Mengklik tombol Edit produk',                  '□ Action',     '-',                           '14'),
    ('Sistem', 'Membuka modal form Edit Produk',               '□ Action',     '-',                           '15'),
    ('Actor',  'Melihat detail produk (read-only)',            '□ Action',     '-',                           '16'),
    ('Actor',  'Selesai',                                      '◉ End Node',   '-',                           '-'),
])

para(
    'Berdasarkan tabel di atas, terdapat dua percabangan yaitu pada langkah 4 '
    '(pemilihan metode penelusuran) dan langkah 12 (pembeda hak akses role). '
    'Admin dapat mengakses fungsi edit produk sedangkan Kasir hanya dapat '
    'melihat informasi produk secara read-only.',
    first_indent=1.25
)

# ============================================================
# 8. ACTIVITY LAPORAN
# ============================================================
heading('8)  Activity Diagram Laporan Penjualan dan Laporan Inventory')
para('Proses pengaksesan dan pengelolaan laporan oleh Admin.', first_indent=1.25)

build_activity_table(8, 'Langkah-langkah Activity Laporan', [
    ('Actor',  'Mulai',                                          '● Start Node', '-',                              '2'),
    ('Actor',  'Memilih jenis laporan',                          '◇ Decision',   'Penjualan → 3 | Inventory → 13', ''),
    ('Actor',  'Membuka halaman Laporan Penjualan (/sales-report/)', '□ Action', '-',                              '4'),
    ('Sistem', 'Menampilkan seluruh data transaksi',             '□ Action',     '-',                              '5'),
    ('Actor',  'Menerapkan filter (tanggal/kasir/metode/merek)', '□ Action',     '-',                              '6'),
    ('Sistem', 'Memfilter dan menghitung total pendapatan',      '□ Action',     '-',                              '7'),
    ('Actor',  'Memilih aksi lanjutan',                          '◇ Decision',   'Detail → 8 | Void → 10 | Cetak → 12', ''),
    ('Actor',  'Mengklik tombol Detail Transaksi',               '□ Action',     '-',                              '9'),
    ('Sistem', 'Menampilkan struk thermal receipt (modal)',       '□ Action',     '-',                              '18'),
    ('Actor',  'Mengklik Batalkan + mengisi alasan void',        '□ Action',     '-',                              '11'),
    ('Sistem', 'Update status=void, kembalikan stok produk',     '□ Action',     '-',                              '18'),
    ('Actor',  'Mengklik tombol Cetak/Ekspor PDF',               '□ Action',     '-',                              '18'),
    ('Sistem', 'Membuka jendela cetak browser (print-ready)',    '□ Action',     '-',                              '18'),
    ('Actor',  'Membuka halaman Laporan Inventory (/inventory-report/)', '□ Action', '-',                          '14'),
    ('Sistem', 'Menampilkan produk dengan stok > 0',             '□ Action',     '-',                              '15'),
    ('Sistem', 'Menghitung total nilai modal & potensi pendapatan','□ Action',   '-',                              '16'),
    ('Actor',  'Menerapkan filter merek/kategori (opsional)',     '□ Action',    '-',                              '17'),
    ('Sistem', 'Memperbarui tampilan & kalkulasi valuasi',        '□ Action',    '-',                              '18'),
    ('Actor',  'Selesai',                                         '◉ End Node',  '-',                              '-'),
])

para(
    'Berdasarkan tabel di atas, terdapat tiga percabangan yaitu pada langkah 2 '
    '(pemilihan jenis laporan), langkah 7 (aksi lanjutan laporan penjualan). '
    'Seluruh jalur bertemu pada langkah 18 sebagai titik akhir proses.',
    first_indent=1.25
)

# ============================================================
# SIMPAN
# ============================================================
output_path = (r'c:\Users\IRVAN SUSANTO\pos_ddshoes2\pos_ddshoes\docs'
               r'\Bab_5_2_1_4_Activity_Tabel.docx')
doc.save(output_path)
print(f'File berhasil dibuat: {output_path}')
