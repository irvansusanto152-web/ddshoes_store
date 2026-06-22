"""
Script untuk membuat dokumen Word sub-bab 5.2.1.2 Perancangan Basis Data
bagian b. Spesifikasi Tabel Database
Berdasarkan 3NF yang sudah disetujui — DD Shoes POS System
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ============================================================
# PENGATURAN HALAMAN (margin 4-3-3-3 standar skripsi Indonesia)
# ============================================================
section = doc.sections[0]
section.page_height = Cm(29.7)
section.page_width  = Cm(21.0)
section.left_margin   = Cm(4.0)
section.right_margin  = Cm(3.0)
section.top_margin    = Cm(3.0)
section.bottom_margin = Cm(3.0)

# ============================================================
# WARNA
# ============================================================
YELLOW_HEADER = RGBColor(0xFF, 0xC0, 0x00)   # kuning emas header tabel
BLACK         = RGBColor(0x00, 0x00, 0x00)
WHITE         = RGBColor(0xFF, 0xFF, 0xFF)

def hex_to_rgb_str(r, g, b):
    return '{:02X}{:02X}{:02X}'.format(r, g, b)

HEADER_FILL = 'FFC000'   # kuning emas
WHITE_FILL  = 'FFFFFF'

# ============================================================
# HELPER: set cell background color
# ============================================================
def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

# ============================================================
# HELPER: set cell borders
# ============================================================
def set_table_borders(table):
    tbl   = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'),   'single')
        border.set(qn('w:sz'),    '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr.append(tblBorders)

# ============================================================
# HELPER: paragraph style
# ============================================================
def set_font(run, bold=False, size=12, italic=False):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic

def add_paragraph(text='', bold=False, size=12,
                  alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  space_before=0, space_after=6,
                  first_indent=None, left_indent=None):
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    pf.line_spacing = Pt(24)
    if first_indent is not None:
        pf.first_line_indent = Cm(first_indent)
    if left_indent is not None:
        pf.left_indent = Cm(left_indent)
    if text:
        run = p.add_run(text)
        set_font(run, bold=bold, size=size)
    return p

def add_heading(text, bold=True, size=12, space_before=14):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(6)
    pf.line_spacing = Pt(24)
    run = p.add_run(text)
    set_font(run, bold=bold, size=size)
    return p

# ============================================================
# HELPER: buat tabel spesifikasi
# ============================================================
def add_spec_table(table_number, table_name_id, table_name_display, columns,
                   intro_text, desc_text):
    """
    columns    = list of tuple: (no, nama_kolom, tipe_data, panjang, null_notnull, default)
    intro_text = kalimat pengantar sebelum tabel
    desc_text  = paragraf penjelasan setelah tabel
    """
    # --- Judul penomoran (sub heading) ---
    p_num = doc.add_paragraph()
    p_num.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_num.paragraph_format.space_before = Pt(12)
    p_num.paragraph_format.space_after  = Pt(4)
    p_num.paragraph_format.line_spacing = Pt(24)
    r = p_num.add_run(f'{table_number}. Tabel {table_name_display}')
    set_font(r, bold=True, size=12)

    # --- Kalimat pengantar ---
    p_intro = doc.add_paragraph()
    p_intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_intro.paragraph_format.space_before = Pt(2)
    p_intro.paragraph_format.space_after  = Pt(4)
    p_intro.paragraph_format.line_spacing = Pt(24)
    p_intro.paragraph_format.first_line_indent = Cm(1.25)
    run_intro = p_intro.add_run(intro_text)
    set_font(run_intro, bold=False, size=12)

    # --- Caption tabel (bold, center, di atas tabel) ---
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_before = Pt(4)
    p_cap.paragraph_format.space_after  = Pt(4)
    p_cap.paragraph_format.line_spacing = Pt(24)
    r_bold = p_cap.add_run(f'Tabel 5.{table_number + 1} ')
    set_font(r_bold, bold=True, size=12)
    r_normal = p_cap.add_run(f'Tabel {table_name_display}')
    set_font(r_normal, bold=False, size=12)

    # --- Buat tabel ---
    col_count = 6
    tbl = doc.add_table(rows=1 + len(columns), cols=col_count)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = 'Table Grid'
    set_table_borders(tbl)

    # Lebar kolom (total ~14cm untuk konten area)
    col_widths = [Cm(1.0), Cm(3.8), Cm(2.5), Cm(1.8), Cm(2.4), Cm(2.5)]
    for row in tbl.rows:
        for i, cell in enumerate(row.cells):
            cell.width = col_widths[i]

    # --- Header row ---
    headers = ['No', 'Nama Kolom', 'Tipe Data', 'Panjang', 'NULL /\nNOT NULL', 'Default']
    header_row = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = header_row.cells[i]
        set_cell_bg(cell, HEADER_FILL)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # Hapus paragraf lama, buat baru
        for para in cell.paragraphs:
            para.clear()
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        run = p.add_run(h)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        run.font.bold = True

    # --- Data rows ---
    for row_idx, col_data in enumerate(columns):
        row = tbl.rows[row_idx + 1]
        for col_idx, val in enumerate(col_data):
            cell = row.cells[col_idx]
            set_cell_bg(cell, WHITE_FILL)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for para in cell.paragraphs:
                para.clear()
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)
            run = p.add_run(str(val))
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            run.font.bold = False

    # --- Paragraf penjelasan setelah tabel ---
    p_desc = doc.add_paragraph()
    p_desc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_desc.paragraph_format.space_before = Pt(4)
    p_desc.paragraph_format.space_after  = Pt(8)
    p_desc.paragraph_format.line_spacing = Pt(24)
    p_desc.paragraph_format.first_line_indent = Cm(1.25)
    run_desc = p_desc.add_run(desc_text)
    set_font(run_desc, bold=False, size=12)


# ============================================================
# JUDUL SUB-BAB
# ============================================================
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
p_title.paragraph_format.space_before = Pt(0)
p_title.paragraph_format.space_after  = Pt(6)
p_title.paragraph_format.line_spacing = Pt(24)
r = p_title.add_run('b.  Spesifikasi Tabel Database')
set_font(r, bold=True, size=12)

add_paragraph(
    'Spesifikasi tabel database menjelaskan struktur detail dari setiap tabel '
    'yang digunakan dalam sistem informasi POS DD Shoes Store. Spesifikasi ini '
    'mencakup nama kolom, tipe data, panjang, constraint NULL/NOT NULL, serta '
    'nilai default masing-masing atribut. Tabel-tabel berikut merupakan hasil '
    'implementasi dari struktur basis data yang telah dinormalisasi hingga '
    'tahap Third Normal Form (3NF).',
    first_indent=1.25
)

# ============================================================
# DEFINISI KOLOM SETIAP TABEL
# Format: (No, Nama Kolom, Tipe Data, Panjang, NULL/NOT NULL, Default)
# ============================================================

# 1. tb_user (Django auth_user)
tbl1_cols = [
    (1,  'user_id',    'INTEGER',  '—',   'NOT NULL', 'Auto Increment (PK)'),
    (2,  'username',   'VARCHAR',  '150', 'NOT NULL', '—'),
    (3,  'password',   'VARCHAR',  '128', 'NOT NULL', '—'),
    (4,  'is_active',  'BOOLEAN',  '—',   'NOT NULL', 'TRUE'),
    (5,  'last_login', 'DATETIME', '—',   'NULL',     'NULL'),
    (6,  'date_joined','DATETIME', '—',   'NOT NULL', 'CURRENT_TIMESTAMP'),
]
add_spec_table(1, 'tb_user', 'User', tbl1_cols,
    intro_text=(
        'Adapun spesifikasi tabel user dapat dilihat pada Tabel 5.2 berikut.'
    ),
    desc_text=(
        'Tabel user merupakan tabel bawaan Django (auth_user) yang digunakan untuk '
        'menyimpan data akun pengguna sistem. Tabel ini memiliki 6 (enam) atribut '
        'utama yaitu user_id sebagai kunci primer dengan nilai auto increment, '
        'username untuk nama login pengguna, password yang disimpan dalam bentuk '
        'hash terenkripsi, is_active untuk menandai status aktif atau nonaktif akun, '
        'last_login untuk mencatat waktu login terakhir, serta date_joined yang '
        'secara otomatis diisi dengan timestamp saat akun dibuat.'
    )
)

# 2. tb_userprofile
tbl2_cols = [
    (1, 'profile_id', 'INTEGER',  '—',  'NOT NULL', 'Auto Increment (PK)'),
    (2, 'user_id',    'INTEGER',  '—',  'NOT NULL', 'FK → tb_user'),
    (3, 'role',       'VARCHAR',  '20', 'NOT NULL', '—'),
    (4, 'phone',      'VARCHAR',  '20', 'NULL',     'NULL'),
]
add_spec_table(2, 'tb_userprofile', 'User Profile', tbl2_cols,
    intro_text=(
        'Adapun spesifikasi tabel user profile dapat dilihat pada Tabel 5.3 berikut.'
    ),
    desc_text=(
        'Tabel user profile merupakan tabel ekstensi dari tabel user yang digunakan '
        'untuk menyimpan informasi tambahan pengguna sistem. Tabel ini memiliki 4 '
        '(empat) atribut yaitu profile_id sebagai kunci primer, user_id sebagai '
        'foreign key yang merujuk ke tabel user dengan relasi one-to-one, role yang '
        'berisi peran pengguna dalam sistem (admin atau kasir), serta phone untuk '
        'menyimpan nomor telepon pengguna yang bersifat opsional.'
    )
)

# 3. tb_kategori
tbl3_cols = [
    (1, 'category_id',   'INTEGER', '—',   'NOT NULL', 'Auto Increment (PK)'),
    (2, 'category_name', 'VARCHAR', '100', 'NOT NULL', '—'),
]
add_spec_table(3, 'tb_kategori', 'Kategori', tbl3_cols,
    intro_text=(
        'Adapun spesifikasi tabel kategori dapat dilihat pada Tabel 5.4 berikut.'
    ),
    desc_text=(
        'Tabel kategori digunakan untuk menyimpan data kategori produk sepatu yang '
        'dijual di DD Shoes Store. Tabel ini memiliki 2 (dua) atribut yaitu '
        'category_id sebagai kunci primer dengan nilai auto increment dan '
        'category_name untuk menyimpan nama kategori sepatu seperti Sneakers, '
        'Boots, Sandal, dan sebagainya.'
    )
)

# 4. tb_merek
tbl4_cols = [
    (1, 'brand_id',   'INTEGER', '—',   'NOT NULL', 'Auto Increment (PK)'),
    (2, 'brand_name', 'VARCHAR', '100', 'NOT NULL', '—'),
]
add_spec_table(4, 'tb_merek', 'Merek', tbl4_cols,
    intro_text=(
        'Adapun spesifikasi tabel merek dapat dilihat pada Tabel 5.5 berikut.'
    ),
    desc_text=(
        'Tabel merek digunakan untuk menyimpan data merek sepatu yang tersedia '
        'dalam katalog DD Shoes Store. Tabel ini memiliki 2 (dua) atribut yaitu '
        'brand_id sebagai kunci primer dengan nilai auto increment dan brand_name '
        'untuk menyimpan nama merek sepatu seperti Nike, Adidas, Vans, '
        'New Balance, dan sebagainya.'
    )
)

# 5. tb_pemasok
tbl5_cols = [
    (1, 'supplier_id',    'INTEGER', '—',   'NOT NULL', 'Auto Increment (PK)'),
    (2, 'supplier_name',  'VARCHAR', '100', 'NOT NULL', '—'),
    (3, 'supplier_phone', 'VARCHAR', '20',  'NULL',     'NULL'),
    (4, 'supplier_notes', 'TEXT',    '—',   'NULL',     'NULL'),
]
add_spec_table(5, 'tb_pemasok', 'Pemasok', tbl5_cols,
    intro_text=(
        'Adapun spesifikasi tabel pemasok dapat dilihat pada Tabel 5.6 berikut.'
    ),
    desc_text=(
        'Tabel pemasok digunakan untuk menyimpan data pihak pemasok atau sumber '
        'barang sepatu yang masuk ke DD Shoes Store. Tabel ini memiliki 4 (empat) '
        'atribut yaitu supplier_id sebagai kunci primer, supplier_name untuk nama '
        'pemasok, supplier_phone untuk nomor telepon pemasok yang bersifat opsional, '
        'serta supplier_notes untuk menyimpan catatan tambahan terkait pemasok '
        'yang juga bersifat opsional.'
    )
)

# 6. tb_produk
tbl6_cols = [
    (1,  'product_id',       'INTEGER',  '—',   'NOT NULL', 'Auto Increment (PK)'),
    (2,  'category_id',      'INTEGER',  '—',   'NULL',     'FK → tb_kategori'),
    (3,  'brand_id',         'INTEGER',  '—',   'NULL',     'FK → tb_merek'),
    (4,  'product_name',     'VARCHAR',  '255', 'NOT NULL', '—'),
    (5,  'size',             'VARCHAR',  '50',  'NOT NULL', '—'),
    (6,  'condition',        'VARCHAR',  '50',  'NOT NULL', '—'),
    (7,  'description',      'TEXT',     '—',   'NULL',     'NULL'),
    (8,  'buy_price',        'INTEGER',  '—',   'NOT NULL', '—'),
    (9,  'sell_price',       'INTEGER',  '—',   'NOT NULL', '—'),
    (10, 'stock',            'INTEGER',  '—',   'NOT NULL', '0'),
    (11, 'image',            'VARCHAR',  '255', 'NULL',     'NULL'),
    (12, 'product_status',   'VARCHAR',  '20',  'NOT NULL', 'active'),
    (13, 'product_created_at','DATETIME','—',   'NOT NULL', 'CURRENT_TIMESTAMP'),
]
add_spec_table(6, 'tb_produk', 'Produk', tbl6_cols,
    intro_text=(
        'Adapun spesifikasi tabel produk dapat dilihat pada Tabel 5.7 berikut.'
    ),
    desc_text=(
        'Tabel produk merupakan tabel sentral dalam sistem yang digunakan untuk '
        'menyimpan seluruh data katalog sepatu DD Shoes Store. Tabel ini memiliki '
        '13 (tiga belas) atribut. product_id merupakan kunci primer, category_id '
        'dan brand_id merupakan foreign key yang merujuk ke tabel kategori dan '
        'tabel merek. Atribut product_name menyimpan nama model sepatu, size untuk '
        'ukuran sepatu, condition untuk grade kondisi barang (Baru, Like New, Good, '
        'atau Fair), description untuk deskripsi detail produk, buy_price dan '
        'sell_price untuk harga beli dan harga jual, stock untuk jumlah stok '
        'tersedia, image untuk path foto produk, product_status untuk status aktif '
        'atau nonaktif produk, serta product_created_at yang otomatis diisi '
        'timestamp saat produk ditambahkan ke sistem.'
    )
)

# 7. tb_barangmasuk
tbl7_cols = [
    (1, 'stockin_id',       'INTEGER',  '—',   'NOT NULL', 'Auto Increment (PK)'),
    (2, 'supplier_id',      'INTEGER',  '—',   'NULL',     'FK → tb_pemasok'),
    (3, 'received_by',      'INTEGER',  '—',   'NULL',     'FK → tb_user'),
    (4, 'received_date',    'DATE',     '—',   'NOT NULL', '—'),
    (5, 'stockin_notes',    'TEXT',     '—',   'NULL',     'NULL'),
    (6, 'stockin_created_at','DATETIME','—',   'NOT NULL', 'CURRENT_TIMESTAMP'),
]
add_spec_table(7, 'tb_barangmasuk', 'Barang Masuk', tbl7_cols,
    intro_text=(
        'Adapun spesifikasi tabel barang masuk dapat dilihat pada Tabel 5.8 berikut.'
    ),
    desc_text=(
        'Tabel barang masuk digunakan untuk menyimpan data header atau dokumen '
        'induk dari setiap sesi penerimaan barang dari pemasok. Tabel ini memiliki '
        '6 (enam) atribut yaitu stockin_id sebagai kunci primer, supplier_id '
        'sebagai foreign key yang merujuk ke tabel pemasok, received_by sebagai '
        'foreign key yang merujuk ke tabel user untuk mencatat admin yang melakukan '
        'pencatatan, received_date untuk tanggal penerimaan barang fisik, '
        'stockin_notes untuk catatan kondisi penerimaan yang bersifat opsional, '
        'serta stockin_created_at yang otomatis diisi timestamp saat data disimpan.'
    )
)

# 8. tb_detailbarangmasuk
tbl8_cols = [
    (1, 'stockin_detail_id', 'INTEGER', '—', 'NOT NULL', 'Auto Increment (PK)'),
    (2, 'stockin_id',        'INTEGER', '—', 'NOT NULL', 'FK → tb_barangmasuk'),
    (3, 'product_id',        'INTEGER', '—', 'NOT NULL', 'FK → tb_produk'),
    (4, 'stockin_qty',       'INTEGER', '—', 'NOT NULL', '—'),
    (5, 'stockin_buy_price', 'INTEGER', '—', 'NOT NULL', '—'),
]
add_spec_table(8, 'tb_detailbarangmasuk', 'Detail Barang Masuk', tbl8_cols,
    intro_text=(
        'Adapun spesifikasi tabel detail barang masuk dapat dilihat pada Tabel 5.9 berikut.'
    ),
    desc_text=(
        'Tabel detail barang masuk digunakan untuk menyimpan data baris item produk '
        'yang diterima dalam satu sesi penerimaan barang. Tabel ini memiliki 5 '
        '(lima) atribut yaitu stockin_detail_id sebagai kunci primer, stockin_id '
        'sebagai foreign key yang merujuk ke tabel barang masuk, product_id sebagai '
        'foreign key yang merujuk ke tabel produk, stockin_qty untuk jumlah unit '
        'produk yang diterima, serta stockin_buy_price untuk harga beli produk pada '
        'saat penerimaan tersebut. Setiap data yang disimpan pada tabel ini akan '
        'secara otomatis menambah nilai stok pada tabel produk melalui mekanisme '
        'Django Signals.'
    )
)

# 9. tb_transaksi
tbl9_cols = [
    (1,  'transaction_id',     'INTEGER',  '—',  'NOT NULL', 'Auto Increment (PK)'),
    (2,  'cashier_id',         'INTEGER',  '—',  'NULL',     'FK → tb_user'),
    (3,  'subtotal_amount',    'INTEGER',  '—',  'NOT NULL', '0'),
    (4,  'discount_amount',    'INTEGER',  '—',  'NOT NULL', '0'),
    (5,  'total_amount',       'INTEGER',  '—',  'NOT NULL', '—'),
    (6,  'payment_method',     'VARCHAR',  '20', 'NOT NULL', '—'),
    (7,  'cash_received',      'INTEGER',  '—',  'NULL',     'NULL'),
    (8,  'change_amount',      'INTEGER',  '—',  'NULL',     'NULL'),
    (9,  'transaction_date',   'DATETIME', '—',  'NOT NULL', 'CURRENT_TIMESTAMP'),
    (10, 'transaction_status', 'VARCHAR',  '20', 'NOT NULL', 'success'),
    (11, 'voided_by',          'INTEGER',  '—',  'NULL',     'FK → tb_user'),
    (12, 'void_reason',        'TEXT',     '—',  'NULL',     'NULL'),
    (13, 'voided_at',          'DATETIME', '—',  'NULL',     'NULL'),
]
add_spec_table(9, 'tb_transaksi', 'Transaksi', tbl9_cols,
    intro_text=(
        'Adapun spesifikasi tabel transaksi dapat dilihat pada Tabel 5.10 berikut.'
    ),
    desc_text=(
        'Tabel transaksi digunakan untuk menyimpan data header atau struk dari '
        'setiap transaksi penjualan yang dilakukan kasir melalui halaman POS. '
        'Tabel ini memiliki 13 (tiga belas) atribut. transaction_id merupakan '
        'kunci primer, cashier_id merupakan foreign key yang merujuk ke tabel user. '
        'Atribut subtotal_amount menyimpan total kotor sebelum diskon, '
        'discount_amount untuk nominal diskon, total_amount untuk total akhir yang '
        'harus dibayar, payment_method untuk metode pembayaran (tunai, transfer_bri, '
        'atau qris), cash_received untuk nominal uang tunai yang diterima kasir, '
        'change_amount untuk kembalian, transaction_date untuk timestamp transaksi, '
        'serta transaction_status untuk status transaksi (success atau void). '
        'Atribut voided_by, void_reason, dan voided_at digunakan untuk mencatat '
        'data pembatalan transaksi apabila dilakukan oleh admin.'
    )
)

# 10. tb_detailtransaksi
tbl10_cols = [
    (1, 'trx_detail_id',  'INTEGER', '—', 'NOT NULL', 'Auto Increment (PK)'),
    (2, 'transaction_id', 'INTEGER', '—', 'NOT NULL', 'FK → tb_transaksi'),
    (3, 'product_id',     'INTEGER', '—', 'NOT NULL', 'FK → tb_produk'),
    (4, 'trx_quantity',   'INTEGER', '—', 'NOT NULL', '—'),
    (5, 'trx_sell_price', 'INTEGER', '—', 'NOT NULL', '—'),
    (6, 'trx_subtotal',   'INTEGER', '—', 'NOT NULL', '—'),
]
add_spec_table(10, 'tb_detailtransaksi', 'Detail Transaksi', tbl10_cols,
    intro_text=(
        'Adapun spesifikasi tabel detail transaksi dapat dilihat pada Tabel 5.11 berikut.'
    ),
    desc_text=(
        'Tabel detail transaksi digunakan untuk menyimpan data baris item produk '
        'yang terjual dalam satu transaksi penjualan. Tabel ini memiliki 6 (enam) '
        'atribut yaitu trx_detail_id sebagai kunci primer, transaction_id sebagai '
        'foreign key yang merujuk ke tabel transaksi, product_id sebagai foreign '
        'key yang merujuk ke tabel produk, trx_quantity untuk jumlah unit produk '
        'yang dibeli, trx_sell_price untuk harga jual produk pada saat transaksi '
        'berlangsung sebagai snapshot harga, serta trx_subtotal yang merupakan '
        'hasil perkalian antara trx_quantity dan trx_sell_price. Setiap data yang '
        'disimpan pada tabel ini akan secara otomatis mengurangi nilai stok pada '
        'tabel produk melalui mekanisme Django Signals.'
    )
)

# 11. tb_closing
tbl11_cols = [
    (1,  'closing_id',            'INTEGER',  '—', 'NOT NULL', 'Auto Increment (PK)'),
    (2,  'cashier_id',            'INTEGER',  '—', 'NOT NULL', 'FK → tb_user'),
    (3,  'closing_date',          'DATE',     '—', 'NOT NULL', '—'),
    (4,  'system_cash_total',     'INTEGER',  '—', 'NOT NULL', '—'),
    (5,  'system_transfer_total', 'INTEGER',  '—', 'NOT NULL', '—'),
    (6,  'system_qris_total',     'INTEGER',  '—', 'NOT NULL', '—'),
    (7,  'actual_cash',           'INTEGER',  '—', 'NOT NULL', '—'),
    (8,  'cash_difference',       'INTEGER',  '—', 'NOT NULL', '—'),
    (9,  'closing_notes',         'TEXT',     '—', 'NULL',     'NULL'),
    (10, 'is_locked',             'BOOLEAN',  '—', 'NOT NULL', 'FALSE'),
    (11, 'submitted_at',          'DATETIME', '—', 'NOT NULL', 'CURRENT_TIMESTAMP'),
    (12, 'unlocked_by',           'INTEGER',  '—', 'NULL',     'FK → tb_user'),
    (13, 'unlocked_at',           'DATETIME', '—', 'NULL',     'NULL'),
    (14, 'unlock_reason',         'TEXT',     '—', 'NULL',     'NULL'),
    (15, 'unlock_count',          'INTEGER',  '—', 'NOT NULL', '0'),
]
add_spec_table(11, 'tb_closing', 'Closing Kasir', tbl11_cols,
    intro_text=(
        'Adapun spesifikasi tabel closing kasir dapat dilihat pada Tabel 5.12 berikut.'
    ),
    desc_text=(
        'Tabel closing kasir digunakan untuk menyimpan data laporan penutupan shift '
        'harian kasir beserta rekonsiliasi jumlah uang kas. Tabel ini memiliki 15 '
        '(lima belas) atribut. closing_id merupakan kunci primer, cashier_id '
        'merupakan foreign key yang merujuk ke tabel user. Atribut closing_date '
        'menyimpan tanggal operasional shift, system_cash_total, '
        'system_transfer_total, dan system_qris_total menyimpan total penjualan '
        'berdasarkan masing-masing metode pembayaran menurut data sistem. '
        'Atribut actual_cash untuk jumlah uang fisik tunai yang dihitung kasir, '
        'cash_difference untuk selisih antara kas fisik dan data sistem, serta '
        'is_locked untuk mengunci data closing agar tidak dapat diubah kembali. '
        'Atribut unlocked_by, unlocked_at, unlock_reason, dan unlock_count '
        'merupakan data audit trail yang dicatat apabila admin membuka kembali '
        'data closing yang sudah dikunci.'
    )
)

# 12. tb_penyesuaianstok
tbl12_cols = [
    (1, 'adjustment_id', 'INTEGER',  '—',  'NOT NULL', 'Auto Increment (PK)'),
    (2, 'product_id',    'INTEGER',  '—',  'NULL',     'FK → tb_produk'),
    (3, 'adjusted_by',   'INTEGER',  '—',  'NULL',     'FK → tb_user'),
    (4, 'adj_quantity',  'INTEGER',  '—',  'NOT NULL', '—'),
    (5, 'adj_reason',    'VARCHAR',  '20', 'NOT NULL', '—'),
    (6, 'adj_notes',     'TEXT',     '—',  'NULL',     'NULL'),
    (7, 'adjusted_at',   'DATETIME', '—',  'NOT NULL', 'CURRENT_TIMESTAMP'),
]
add_spec_table(12, 'tb_penyesuaianstok', 'Penyesuaian Stok', tbl12_cols,
    intro_text=(
        'Adapun spesifikasi tabel penyesuaian stok dapat dilihat pada Tabel 5.13 berikut.'
    ),
    desc_text=(
        'Tabel penyesuaian stok digunakan untuk mencatat setiap perubahan atau '
        'pengurangan stok yang tidak berasal dari jalur penjualan maupun '
        'penerimaan barang, misalnya karena barang rusak, hilang, atau dikembalikan '
        'ke pemasok. Tabel ini memiliki 7 (tujuh) atribut yaitu adjustment_id '
        'sebagai kunci primer, product_id sebagai foreign key yang merujuk ke tabel '
        'produk, adjusted_by sebagai foreign key yang merujuk ke tabel user untuk '
        'mencatat admin yang melakukan koreksi, adj_quantity untuk jumlah koreksi '
        'stok yang disimpan dalam nilai negatif, adj_reason untuk kategori penyebab '
        'koreksi (rusak, hilang, retur, atau lainnya), adj_notes untuk catatan '
        'detail yang bersifat opsional, serta adjusted_at yang otomatis diisi '
        'timestamp saat penyesuaian dilakukan.'
    )
)

# ============================================================
# SIMPAN FILE
# ============================================================
output_path = r'c:\Users\IRVAN SUSANTO\pos_ddshoes2\pos_ddshoes\docs\Bab_5_2_1_2_Spesifikasi_Tabel_v2.docx'
doc.save(output_path)
print(f'File berhasil dibuat: {output_path}')
