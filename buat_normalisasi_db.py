import os

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    import sys
    print("Error: Pustaka python-docx belum terinstal.")
    print("Jalankan: pip install python-docx")
    sys.exit(1)


# ─────────────────────────────────────────────
# HELPER FUNCTIONS
# ─────────────────────────────────────────────

def add_justified_paragraph(doc, text, space_before=0, space_after=6, first_line_indent=True):
    """Tambah paragraf rata kanan-kiri dengan indentasi baris pertama."""
    p = doc.add_paragraph(text)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if first_line_indent:
        pf.first_line_indent = Cm(1.27)  # 0.5 inch
    return p


def set_table_style(table):
    """Terapkan style tabel akademis sederhana (border hitam tipis)."""
    table.style = 'Table Grid'
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for para in cell.paragraphs:
                para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                para.paragraph_format.space_before = Pt(2)
                para.paragraph_format.space_after = Pt(2)


def shade_cell(cell, fill_hex):
    """Warnai latar belakang sel tabel."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)


def add_table_with_header(doc, headers, rows_data, header_fill='4472C4', col_widths=None):
    """
    Buat tabel dengan header berwarna dan data baris.
    headers: list of str
    rows_data: list of list of str
    """
    table = doc.add_table(rows=1 + len(rows_data), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_style(table)

    # Header row
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        shade_cell(cell, header_fill)
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)

    # Data rows — alternating row color
    for ri, row_data in enumerate(rows_data):
        row = table.rows[ri + 1]
        fill = 'DEEAF1' if ri % 2 == 0 else 'FFFFFF'
        for ci, cell_text in enumerate(row_data):
            cell = row.cells[ci]
            shade_cell(cell, fill)
            run = cell.paragraphs[0].add_run(str(cell_text))
            run.font.size = Pt(9)

    # Set column widths if provided
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    return table


def add_caption(doc, text):
    """Tambah keterangan tabel/gambar di bawah."""
    p = doc.add_paragraph(text)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.runs[0].italic = True
    p.runs[0].font.size = Pt(10)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(10)


# ─────────────────────────────────────────────
# DOKUMEN UTAMA
# ─────────────────────────────────────────────

doc = Document()

# Margin halaman (2.5 cm semua sisi — APA style umum)
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(4.0)   # margin kiri skripsi biasa 4 cm
    section.right_margin  = Cm(3.0)

# ============================================================
# JUDUL BAGIAN
# ============================================================
h = doc.add_heading('5.2.1.2 Perancangan Basis Data', level=3)
h.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

h2 = doc.add_heading('a. Normalisasi Basis Data', level=4)
h2.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

# ============================================================
# PENGANTAR NORMALISASI
# ============================================================
add_justified_paragraph(doc,
    "Normalisasi merupakan proses sistematis dalam perancangan basis data yang bertujuan untuk "
    "mengurangi redundansi data serta mencegah terjadinya anomali pada saat operasi penyisipan "
    "(insert), pembaruan (update), maupun penghapusan (delete) data. Proses normalisasi dilakukan "
    "secara bertahap mulai dari bentuk tidak normal (Unnormalized Form/UNF) hingga mencapai "
    "bentuk normal ketiga (Third Normal Form/3NF), sesuai dengan referensi Database System "
    "Concepts (Silberschatz, Korth dan Sudarshan)."
)

add_justified_paragraph(doc,
    "Perancangan basis data sistem informasi POS DD Shoes Store didasarkan pada dokumen-dokumen "
    "bisnis yang digunakan dalam operasional toko sehari-hari, meliputi struk penjualan, formulir "
    "penerimaan barang masuk, data katalog produk, data akun pengguna, serta formulir penutupan "
    "kasir (closing shift). Dokumen-dokumen tersebut menjadi sumber identifikasi atribut pada "
    "tahap awal normalisasi."
)

# ============================================================
# UNF
# ============================================================
doc.add_heading('1) Unnormalized Form (UNF)', level=5)

add_justified_paragraph(doc,
    "Unnormalized Form (UNF) merupakan bentuk awal representasi data yang diperoleh langsung dari "
    "dokumen atau formulir bisnis yang digunakan pada UMKM DD Shoes Store. Pada tahap ini, seluruh "
    "atribut dimasukkan dalam satu entitas tanpa mewajibkan adanya format tertentu. UNF "
    "memungkinkan adanya data berulang (repeating groups) yang dapat menjadi sumber anomali data "
    "apabila tidak dinormalisasi lebih lanjut."
)

add_justified_paragraph(doc,
    "Berdasarkan identifikasi terhadap dokumen bisnis yang relevan, terdapat lima entitas utama "
    "yang teridentifikasi pada tahap UNF, yaitu: Penjualan, Barang Masuk, Produk, Pengguna, "
    "dan Penutupan Kasir (Closing). Berikut adalah representasi masing-masing entitas dalam "
    "bentuk UNF:"
)

# --- Tabel UNF 1: PENJUALAN ---
doc.add_paragraph("Tabel UNF 1: Entitas PENJUALAN (dari Struk/Nota Penjualan)",
                  style='Intense Quote')

unf_headers = ['No', 'Atribut', 'Keterangan']
unf1_headers = unf_headers
unf1_data = [
    ['1',  'No_Struk',           'Nomor struk transaksi'],
    ['2',  'Tanggal',            'Tanggal transaksi'],
    ['3',  'Nama_Kasir',         'Nama kasir yang melayani'],
    ['4',  'Metode_Bayar',       'Cara pembayaran (Tunai/Transfer/QRIS)'],
    ['5',  'Total',              'Total nilai transaksi'],
    ['6',  'Kas_Diterima',       'Uang yang diterima dari pelanggan'],
    ['7',  'Kembalian',          'Uang kembalian untuk pelanggan'],
    ['8',  '* Nama_Produk',      'Nama produk yang dibeli [REPEATING]'],
    ['9',  '* Merek',            'Merek produk yang dibeli [REPEATING]'],
    ['10', '* Ukuran',           'Ukuran sepatu yang dibeli [REPEATING]'],
    ['11', '* Kondisi',          'Kondisi produk yang dibeli [REPEATING]'],
    ['12', '* Harga_Jual',       'Harga satuan produk [REPEATING]'],
    ['13', '* Qty',              'Jumlah unit yang dibeli [REPEATING]'],
    ['14', '* Subtotal',         'Qty × Harga_Jual [REPEATING]'],
]
add_table_with_header(doc, unf1_headers, unf1_data, col_widths=[1, 5, 8])
add_caption(doc, "Keterangan: * = Atribut yang bersifat repeating group")

# --- Tabel UNF 2: BARANG_MASUK ---
doc.add_paragraph("Tabel UNF 2: Entitas BARANG_MASUK (dari Formulir Penerimaan Barang)",
                  style='Intense Quote')

unf2_headers = unf_headers
unf2_data = [
    ['1',  'No_Dokumen',         'Nomor dokumen penerimaan'],
    ['2',  'Tanggal_Terima',     'Tanggal barang diterima'],
    ['3',  'Nama_Pemasok',       'Nama pemasok/sumber barang'],
    ['4',  'No_HP_Pemasok',      'Nomor HP pemasok'],
    ['5',  'Catatan',            'Catatan kondisi kiriman'],
    ['6',  'Dicatat_Oleh',       'Admin yang mencatat penerimaan'],
    ['7',  '* Nama_Produk',      'Nama produk yang masuk [REPEATING]'],
    ['8',  '* Merek',            'Merek produk yang masuk [REPEATING]'],
    ['9',  '* Kategori',         'Kategori produk [REPEATING]'],
    ['10', '* Ukuran',           'Ukuran sepatu [REPEATING]'],
    ['11', '* Kondisi',          'Kondisi produk [REPEATING]'],
    ['12', '* Harga_Beli',       'Harga beli dari pemasok [REPEATING]'],
    ['13', '* Qty',              'Jumlah unit yang masuk [REPEATING]'],
]
add_table_with_header(doc, unf2_headers, unf2_data, col_widths=[1, 5, 8])
add_caption(doc, "Keterangan: * = Atribut yang bersifat repeating group")

# --- Tabel UNF 3: PRODUK ---
unf3_headers = unf_headers
doc.add_paragraph("Tabel UNF 3: Entitas PRODUK (dari Katalog Toko)",
                  style='Intense Quote')

unf3_data = [
    ['1', 'Nama_Produk',   'Nama produk sepatu'],
    ['2', 'Merek',         'Merek sepatu (Nike, Adidas, dll)'],
    ['3', 'Kategori',      'Kategori sepatu (Sneakers, Boots, dll)'],
    ['4', 'Ukuran',        'Nomor ukuran sepatu'],
    ['5', 'Kondisi',       'Grade kondisi (Baru/Like New/Good/Fair)'],
    ['6', 'Deskripsi',     'Keterangan detail produk'],
    ['7', 'Harga_Beli',    'Harga beli dari pemasok'],
    ['8', 'Harga_Jual',    'Harga jual ke pelanggan'],
    ['9', 'Stok',          'Jumlah unit tersedia'],
    ['10','Foto',          'Path foto produk'],
    ['11','Status',        'Status produk (aktif/nonaktif)'],
]
add_table_with_header(doc, unf3_headers, unf3_data, col_widths=[1, 5, 8])
add_caption(doc, "Tidak terdapat repeating group pada entitas PRODUK")

# --- Tabel UNF 4: PENGGUNA ---
doc.add_paragraph("Tabel UNF 4: Entitas PENGGUNA (Data Akun Admin dan Kasir)",
                  style='Intense Quote')

unf4_data = [
    ['1', 'Username',       'Nama login pengguna'],
    ['2', 'Password',       'Kata sandi (terenkripsi)'],
    ['3', 'Role',           'Peran pengguna (admin/kasir)'],
    ['4', 'No_HP',          'Nomor telepon pengguna'],
    ['5', 'Status_Aktif',   'Status keaktifan akun'],
    ['6', 'Waktu_Login',    'Waktu login terakhir'],
]
add_table_with_header(doc, unf3_headers, unf4_data, col_widths=[1, 5, 8])
add_caption(doc, "Tidak terdapat repeating group pada entitas PENGGUNA")

# --- Tabel UNF 5: CLOSING ---
doc.add_paragraph("Tabel UNF 5: Entitas PENUTUPAN_KASIR (dari Formulir Closing Shift)",
                  style='Intense Quote')

unf5_data = [
    ['1', 'Nama_Kasir',           'Kasir yang melakukan closing'],
    ['2', 'Tanggal_Shift',        'Tanggal shift berlangsung'],
    ['3', 'Total_Tunai_Sistem',   'Total tunai menurut catatan sistem'],
    ['4', 'Total_Transfer_Sistem','Total transfer menurut sistem'],
    ['5', 'Total_QRIS_Sistem',    'Total QRIS menurut sistem'],
    ['6', 'Kas_Fisik_Aktual',     'Uang tunai fisik yang dihitung kasir'],
    ['7', 'Selisih',              'Selisih (Kas_Fisik - Total_Tunai)'],
    ['8', 'Catatan',              'Catatan kasir (opsional)'],
    ['9', 'Status_Kunci',         'True jika sudah disubmit (locked)'],
]
add_table_with_header(doc, unf3_headers, unf5_data, col_widths=[1, 5, 8])
add_caption(doc, "Tidak terdapat repeating group pada entitas PENUTUPAN_KASIR")

# ============================================================
# 1NF
# ============================================================
doc.add_heading('2) First Normal Form (1NF)', level=5)

add_justified_paragraph(doc,
    "First Normal Form (1NF) mensyaratkan bahwa setiap kolom dalam tabel harus memiliki nilai "
    "yang unik dan atomik, serta tidak boleh terdapat kelompok data yang berulang (repeating "
    "groups) dalam satu baris. Pada tahap ini dilakukan pengelompokkan beberapa tipe data yang "
    "sejenis ke dalam tabel terpisah sehingga dapat mengatasi anomali data."
)

add_justified_paragraph(doc,
    "Berdasarkan analisis UNF, terdapat dua entitas yang memiliki repeating groups, yaitu "
    "PENJUALAN dan BARANG_MASUK. Kedua entitas tersebut dipecah menjadi tabel header dan tabel "
    "detail. Seluruh tabel juga dilengkapi dengan Primary Key (PK) yang bersifat unik. "
    "Hasil transformasi ke 1NF menghasilkan delapan tabel sebagai berikut:"
)

# Penjelasan transformasi
add_justified_paragraph(doc,
    "Transformasi yang dilakukan dari UNF ke 1NF adalah sebagai berikut: entitas PENJUALAN "
    "yang memiliki repeating group produk-produk yang dibeli, dipecah menjadi dua tabel yaitu "
    "TRANSAKSI (menyimpan data header transaksi) dan DETAIL_TRANSAKSI (menyimpan data setiap "
    "item produk yang terjual). Demikian pula entitas BARANG_MASUK dipecah menjadi tabel "
    "STOCK_IN (header penerimaan) dan DETAIL_STOCK_IN (baris item produk yang masuk). "
    "Sementara tiga entitas lainnya yaitu PRODUK, PENGGUNA, dan PENUTUPAN_KASIR tidak mengalami "
    "pemecahan karena tidak memiliki repeating group, namun tetap ditambahkan Primary Key."
)

# Tabel ringkasan 1NF
doc.add_paragraph("Tabel 1NF: Hasil Transformasi ke First Normal Form",
                  style='Intense Quote')

onf_headers = ['No', 'Nama Tabel', 'Primary Key', 'Atribut']
onf_data = [
    ['1', 'TRANSAKSI',
     'id_transaksi (PK)',
     'id_transaksi, tanggal, id_kasir, metode_bayar, total, kas_diterima, kembalian'],
    ['2', 'DETAIL_TRANSAKSI',
     'id_detail (PK)',
     'id_detail, id_transaksi, nama_produk, merek, ukuran, kondisi, harga_jual, qty, subtotal'],
    ['3', 'STOCK_IN',
     'id_stock_in (PK)',
     'id_stock_in, tanggal_terima, nama_pemasok, no_hp_pemasok, catatan, dicatat_oleh'],
    ['4', 'DETAIL_STOCK_IN',
     'id_detail_si (PK)',
     'id_detail_si, id_stock_in, nama_produk, merek, kategori, ukuran, kondisi, harga_beli, qty'],
    ['5', 'PRODUK',
     'id_produk (PK)',
     'id_produk, nama, merek, kategori, ukuran, kondisi, deskripsi, harga_beli, harga_jual, stok, foto, status'],
    ['6', 'PENGGUNA',
     'id_pengguna (PK)',
     'id_pengguna, username, password, role, no_hp, status_aktif, waktu_login'],
    ['7', 'PEMASOK',
     'id_pemasok (PK)',
     'id_pemasok, nama_pemasok, no_hp, catatan'],
    ['8', 'PENUTUPAN_KASIR',
     'id_closing (PK)',
     'id_closing, id_kasir, tanggal_shift, total_tunai_sistem, total_transfer_sistem, total_qris_sistem, kas_fisik_aktual, selisih, catatan, status_kunci'],
]
add_table_with_header(doc, onf_headers, onf_data, col_widths=[0.8, 3.5, 4, 6])
add_caption(doc, "Tabel 1NF — Semua repeating group telah dieliminasi dan setiap tabel memiliki Primary Key")

# ============================================================
# 2NF
# ============================================================
doc.add_heading('3) Second Normal Form (2NF)', level=5)

add_justified_paragraph(doc,
    "Second Normal Form (2NF) mensyaratkan bahwa setiap kolom dalam tabel hanya bergantung "
    "pada satu kunci primer (Primary Key) secara penuh dan tidak bergantung secara parsial "
    "(partial dependency) pada bagian dari composite key. Tahap 2NF juga mewajibkan setiap "
    "tabel yang sudah dipisahkan memiliki kunci primer tersendiri."
)

add_justified_paragraph(doc,
    "Analisis partial dependency dilakukan pada tabel DETAIL_TRANSAKSI dan DETAIL_STOCK_IN "
    "yang berpotensi menyimpan atribut yang seharusnya bergantung pada entitas lain. "
    "Pada tabel DETAIL_TRANSAKSI, atribut nama_produk, merek, ukuran, kondisi, dan harga_jual "
    "sejatinya bergantung pada identitas produk, bukan pada id_transaksi. Dengan demikian, "
    "atribut-atribut tersebut hanya perlu direferensikan melalui foreign key id_produk ke "
    "tabel PRODUK. Demikian pula pada tabel DETAIL_STOCK_IN, atribut nama_produk, merek, "
    "kategori, ukuran, dan kondisi dieliminasi karena bergantung pada id_produk."
)

add_justified_paragraph(doc,
    "Selain itu, pada tabel STOCK_IN, atribut nama_pemasok dan no_hp_pemasok tidak bergantung "
    "pada id_stock_in melainkan pada identitas pemasok. Atribut-atribut tersebut dipindahkan "
    "ke tabel PEMASOK yang telah diidentifikasi sebelumnya, dan pada tabel STOCK_IN hanya "
    "disimpan referensi id_pemasok sebagai foreign key. Demikian pula pada tabel TRANSAKSI, "
    "nama_kasir diganti dengan referensi id_kasir ke tabel PENGGUNA."
)

doc.add_paragraph("Tabel 2NF: Hasil Transformasi ke Second Normal Form",
                  style='Intense Quote')

twonf_headers = ['No', 'Nama Tabel', 'Primary Key', 'Atribut (setelah eliminasi partial dependency)']
twonf_data = [
    ['1', 'TRANSAKSI',
     'id_transaksi (PK)',
     'id_transaksi, tanggal, id_kasir (FK→PENGGUNA), metode_bayar, total, kas_diterima, kembalian'],
    ['2', 'DETAIL_TRANSAKSI',
     'id_detail (PK)',
     'id_detail, id_transaksi (FK), id_produk (FK→PRODUK), harga_jual, qty, subtotal'],
    ['3', 'STOCK_IN',
     'id_stock_in (PK)',
     'id_stock_in, tanggal_terima, id_pemasok (FK→PEMASOK), catatan, id_admin (FK→PENGGUNA)'],
    ['4', 'DETAIL_STOCK_IN',
     'id_detail_si (PK)',
     'id_detail_si, id_stock_in (FK), id_produk (FK→PRODUK), harga_beli, qty'],
    ['5', 'PRODUK',
     'id_produk (PK)',
     'id_produk, nama, merek, kategori, ukuran, kondisi, deskripsi, harga_beli, harga_jual, stok, foto, status'],
    ['6', 'PENGGUNA',
     'id_pengguna (PK)',
     'id_pengguna, username, password, role, no_hp, status_aktif, waktu_login'],
    ['7', 'PEMASOK',
     'id_pemasok (PK)',
     'id_pemasok, nama_pemasok, no_hp, catatan'],
    ['8', 'PENUTUPAN_KASIR',
     'id_closing (PK)',
     'id_closing, id_kasir (FK→PENGGUNA), tanggal_shift, total_tunai_sistem, total_transfer_sistem, total_qris_sistem, kas_fisik_aktual, selisih, catatan, status_kunci'],
]
add_table_with_header(doc, twonf_headers, twonf_data, col_widths=[0.8, 3.5, 4, 7])
add_caption(doc, "Tabel 2NF — Partial dependency telah dieliminasi; atribut yang bergantung pada entitas lain direferensikan melalui FK")

# ============================================================
# 3NF
# ============================================================
doc.add_heading('4) Third Normal Form (3NF)', level=5)

add_justified_paragraph(doc,
    "Third Normal Form (3NF) mensyaratkan bahwa setiap kolom dalam tabel hanya bergantung "
    "pada satu kunci primer dan tidak bergantung pada kolom lain yang bukan kunci dalam tabel "
    "(transitive dependency). 3NF memisahkan atribut yang tidak bergantung langsung dengan "
    "kunci primer, tetapi bergantung pada atribut non-key lainnya."
)

add_justified_paragraph(doc,
    "Pada hasil 2NF, tabel PRODUK masih menyimpan atribut merek dan kategori sebagai nilai "
    "teks secara langsung. Kondisi ini menciptakan ketergantungan transitif: "
    "id_produk → merek → (data merek lainnya), dan id_produk → kategori → (data kategori lainnya). "
    "Artinya, apabila nama merek atau kategori berubah, perubahan harus dilakukan pada seluruh "
    "baris produk yang memiliki merek atau kategori tersebut — kondisi ini menimbulkan update "
    "anomaly. Untuk mengatasi hal tersebut, atribut merek dan kategori dipisahkan menjadi "
    "tabel MEREK dan tabel KATEGORI tersendiri, dan tabel PRODUK hanya menyimpan "
    "id_merek serta id_kategori sebagai foreign key."
)

add_justified_paragraph(doc,
    "Setelah proses normalisasi hingga 3NF selesai, diperoleh sepuluh tabel yang bersih dari "
    "redundansi dan anomali data. Struktur basis data akhir ini kemudian diimplementasikan ke "
    "dalam sistem menggunakan Django ORM dengan basis data SQLite3."
)

doc.add_paragraph("Tabel 3NF: Hasil Akhir Normalisasi — Third Normal Form",
                  style='Intense Quote')

threenf_headers = ['No', 'Nama Tabel (Model Django)', 'Primary Key', 'Atribut Final']
threenf_data = [
    ['1',  'Transactions\n(TRANSAKSI)',
     'id (PK, auto)',
     'id, cashier_id (FK→User), total_amount, payment_method,\ncash_received, change_amount, transaction_date'],
    ['2',  'TransactionDetails\n(DETAIL_TRANSAKSI)',
     'id (PK, auto)',
     'id, transaction_id (FK), product_id (FK→Products),\nquantity, sell_price, subtotal'],
    ['3',  'StockIns\n(STOCK_IN)',
     'id (PK, auto)',
     'id, supplier_id (FK→Suppliers), received_by (FK→User),\nreceived_date, notes, created_at'],
    ['4',  'StockInDetails\n(DETAIL_STOCK_IN)',
     'id (PK, auto)',
     'id, stock_in_id (FK), product_id (FK→Products),\nquantity, buy_price'],
    ['5',  'Products\n(PRODUK)',
     'id (PK, auto)',
     'id, brand_id (FK→Brands), category_id (FK→Categories),\nname, size, condition, description,\nbuy_price, sell_price, stock, image, status, created_at'],
    ['6',  'Brands\n(MEREK)',
     'id (PK, auto)',
     'id, name'],
    ['7',  'Categories\n(KATEGORI)',
     'id (PK, auto)',
     'id, name'],
    ['8',  'Suppliers\n(PEMASOK)',
     'id (PK, auto)',
     'id, name, phone, notes'],
    ['9',  'UserProfile\n(PENGGUNA — Ext. Django)',
     'id (PK, auto)',
     'id, user_id (FK→auth_user, OneToOne),\nrole, phone'],
    ['10', 'CashClosings\n(PENUTUPAN_KASIR)',
     'id (PK, auto)',
     'id, cashier_id (FK→User), closing_date,\nsystem_cash_total, system_transfer_total, system_qris_total,\nactual_cash, cash_difference, notes, is_locked, submitted_at'],
]
add_table_with_header(doc, threenf_headers, threenf_data, col_widths=[0.8, 4, 4, 7.5])
add_caption(doc,
    "Tabel 3NF Final — Sepuluh tabel yang bebas dari partial dependency dan transitive dependency. "
    "Nama dalam kurung merupakan nama model Django yang diimplementasikan dalam sistem."
)

add_justified_paragraph(doc,
    "Berdasarkan hasil normalisasi yang telah dilakukan, dapat disimpulkan bahwa proses "
    "normalisasi dari UNF hingga 3NF menghasilkan sepuluh tabel relasional yang terstruktur "
    "dengan baik. Setiap tabel memiliki satu Primary Key, seluruh atribut bergantung penuh "
    "pada Primary Key tanpa adanya partial dependency maupun transitive dependency. "
    "Struktur basis data ini kemudian diimplementasikan ke dalam sistem POS DD Shoes Store "
    "menggunakan Django ORM dengan penyimpanan pada basis data SQLite3, sebagaimana "
    "didefinisikan dalam berkas models.py pada aplikasi core."
)

# ============================================================
# SIMPAN FILE
# ============================================================
save_dir = r"c:\Users\IRVAN SUSANTO\pos_ddshoes2\pos_ddshoes\docs"
os.makedirs(save_dir, exist_ok=True)
file_path = os.path.join(save_dir, "Bab_5_2_1_2_Normalisasi_Basis_Data.docx")
doc.save(file_path)
print("Sukses! File berhasil dibuat di:")
print(f"   {file_path}")
