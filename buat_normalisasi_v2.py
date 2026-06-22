import os

try:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
except ImportError:
    import sys
    print("Error: Pustaka python-docx belum terinstal.")
    sys.exit(1)

def add_justified_paragraph(doc, text, space_before=0, space_after=6, first_line_indent=True):
    p = doc.add_paragraph(text)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if first_line_indent:
        pf.first_line_indent = Cm(1.27)
    return p

doc = Document()

# Margin
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(4.0)
    section.right_margin  = Cm(3.0)

h = doc.add_heading('5.2.1.2 Perancangan Basis Data', level=3)
h.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

h2 = doc.add_heading('a. Normalisasi Basis Data', level=4)
h2.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

add_justified_paragraph(doc, "Normalisasi merupakan proses pengelompokan atribut data yang membentuk entitas yang sederhana, fleksibel, non redundan serta mudah beradaptasi. Berikut adalah tahapan normalisasi basis data untuk sistem informasi POS dan Inventory DD Shoes Store.")

# --- UNF ---
p_unf = doc.add_paragraph()
p_unf.add_run("UNF (Bentuk Tidak Normal)\n").bold = True
unf_text = (
    "username + password + role + phone_number + is_active + last_login + "
    "supplier_name + supplier_phone + supplier_notes + category_name + brand_name + "
    "product_name + size + condition + description + buy_price + sell_price + stock + "
    "image + status + created_at + received_date + notes + quantity + stockin_buy_price + "
    "subtotal_amount + discount_amount + total_amount + payment_method + cash_received + "
    "change_amount + transaction_date + transaction_status + voided_by + void_reason + voided_at + "
    "transaction_qty + transaction_sell_price + subtotal + closing_date + "
    "system_cash_total + system_transfer_total + system_qris_total + actual_cash + "
    "cash_difference + closing_notes + is_locked + submitted_at + unlocked_by + unlocked_at + "
    "unlock_reason + unlock_count + adjusted_by + adjust_quantity + adjust_reason + adjust_notes + adjusted_at."
)
p_unf.add_run(unf_text)
p_unf.paragraph_format.space_after = Pt(12)

# --- 1NF ---
p_1nf = doc.add_paragraph()
p_1nf.add_run("1NF\n").bold = True
nf1_lines = [
    "tb_pengguna \t= user_id + username + password + role + phone_number + is_active + last_login.",
    "tb_produk \t= product_id + product_name + category_name + brand_name + size + condition + description + buy_price + sell_price + stock + image + status + created_at.",
    "tb_barang_masuk \t= stockin_id + received_date + supplier_name + supplier_phone + supplier_notes + username + notes + created_at.",
    "tb_detail_barang_masuk = detail_stockin_id + stockin_id + product_id + product_name + quantity + buy_price.",
    "tb_transaksi \t= transaction_id + username + subtotal_amount + discount_amount + total_amount + payment_method + cash_received + change_amount + transaction_date + transaction_status + voided_by + void_reason + voided_at.",
    "tb_detail_transaksi \t= detail_transaction_id + transaction_id + product_id + product_name + quantity + sell_price + subtotal.",
    "tb_penutupan_kasir \t= closing_id + username + closing_date + system_cash_total + system_transfer_total + system_qris_total + actual_cash + cash_difference + closing_notes + is_locked + submitted_at + unlocked_by + unlocked_at + unlock_reason + unlock_count.",
    "tb_penyesuaian_stok \t= adjustment_id + product_id + product_name + username + adjust_quantity + adjust_reason + adjust_notes + adjusted_at."
]
for line in nf1_lines:
    p_1nf.add_run(line + "\n")
p_1nf.paragraph_format.space_after = Pt(12)

# --- 2NF ---
p_2nf = doc.add_paragraph()
p_2nf.add_run("2NF\n").bold = True
nf2_lines = [
    "tb_pengguna \t= user_id + username + password + role + phone_number + is_active + last_login.",
    "tb_kategori \t= category_id + category_name.",
    "tb_merek \t= brand_id + brand_name.",
    "tb_pemasok \t= supplier_id + supplier_name + supplier_phone + supplier_notes.",
    "tb_produk \t= product_id + category_id + brand_id + product_name + size + condition + description + buy_price + sell_price + stock + image + status + created_at.",
    "tb_barang_masuk \t= stockin_id + supplier_id + user_id + received_date + notes + created_at.",
    "tb_detail_barang_masuk = detail_stockin_id + stockin_id + product_id + quantity + buy_price.",
    "tb_transaksi \t= transaction_id + user_id + subtotal_amount + discount_amount + total_amount + payment_method + cash_received + change_amount + transaction_date + transaction_status + voided_by + void_reason + voided_at.",
    "tb_detail_transaksi \t= detail_transaction_id + transaction_id + product_id + quantity + sell_price + subtotal.",
    "tb_penutupan_kasir \t= closing_id + user_id + closing_date + system_cash_total + system_transfer_total + system_qris_total + actual_cash + cash_difference + closing_notes + is_locked + submitted_at + unlocked_by + unlocked_at + unlock_reason + unlock_count.",
    "tb_penyesuaian_stok \t= adjustment_id + product_id + user_id + adjust_quantity + adjust_reason + adjust_notes + adjusted_at."
]
for line in nf2_lines:
    p_2nf.add_run(line + "\n")
p_2nf.paragraph_format.space_after = Pt(12)

# --- 3NF ---
p_3nf = doc.add_paragraph()
p_3nf.add_run("3NF\n").bold = True
nf3_lines = [
    "tb_pengguna \t= user_id [PK] + username + password + role + phone_number + is_active + last_login.",
    "tb_kategori \t= category_id [PK] + category_name.",
    "tb_merek \t= brand_id [PK] + brand_name.",
    "tb_pemasok \t= supplier_id [PK] + supplier_name + supplier_phone + supplier_notes.",
    "tb_produk \t= product_id [PK] + category_id [FK] + brand_id [FK] + product_name + size + condition + description + buy_price + sell_price + stock + image + status + created_at.",
    "tb_barang_masuk \t= stockin_id [PK] + supplier_id [FK] + user_id [FK] + received_date + notes + created_at.",
    "tb_detail_barang_masuk = detail_stockin_id [PK] + stockin_id [FK] + product_id [FK] + quantity + buy_price.",
    "tb_transaksi \t= transaction_id [PK] + user_id [FK] + subtotal_amount + discount_amount + total_amount + payment_method + cash_received + change_amount + transaction_date + transaction_status + voided_by [FK] + void_reason + voided_at.",
    "tb_detail_transaksi \t= detail_transaction_id [PK] + transaction_id [FK] + product_id [FK] + quantity + sell_price + subtotal.",
    "tb_penutupan_kasir \t= closing_id [PK] + user_id [FK] + closing_date + system_cash_total + system_transfer_total + system_qris_total + actual_cash + cash_difference + closing_notes + is_locked + submitted_at + unlocked_by [FK] + unlocked_at + unlock_reason + unlock_count.",
    "tb_penyesuaian_stok \t= adjustment_id [PK] + product_id [FK] + user_id [FK] + adjust_quantity + adjust_reason + adjust_notes + adjusted_at."
]
for line in nf3_lines:
    p_3nf.add_run(line + "\n")

# Save
save_dir = r"c:\Users\IRVAN SUSANTO\pos_ddshoes2\pos_ddshoes\docs"
os.makedirs(save_dir, exist_ok=True)
file_path = os.path.join(save_dir, "Bab_5_2_1_2_Normalisasi_Versi_Text_Rev2.docx")
doc.save(file_path)
print("Sukses! File berhasil dibuat di:")
print(f"   {file_path}")
