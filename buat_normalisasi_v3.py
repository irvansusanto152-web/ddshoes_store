"""
Script untuk membuat dokumen Word sub-bab 5.2.1.2 Perancangan Basis Data
bagian a. Normalisasi (UNF, 1NF, 2NF, 3NF)
Berdasarkan database project DD Shoes POS System
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ============================================================
# PENGATURAN HALAMAN (margin 4-3-3-3 standar skripsi Indonesia)
# ============================================================
section = doc.sections[0]
section.page_height = Cm(29.7)
section.page_width  = Cm(21.0)
section.left_margin   = Cm(4.0)
section.right_margin  = Cm(3.0)
section.top_margin    = Cm(3.0)
section.bottom_margin = Cm(3.0)

# ============================================================
# STYLE HELPER
# ============================================================
def set_font(run, bold=False, size=12, color=None):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_paragraph(text='', bold=False, size=12, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  space_before=0, space_after=6, first_indent=None):
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    pf.line_spacing = Pt(24)  # spasi 2
    if first_indent is not None:
        pf.first_line_indent = Cm(first_indent)
    if text:
        run = p.add_run(text)
        set_font(run, bold=bold, size=size)
    return p

def add_heading(text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after  = Pt(6)
    pf.line_spacing = Pt(24)
    run = p.add_run(text)
    set_font(run, bold=True, size=12)
    return p

def add_formula_paragraph(text):
    """Untuk baris formula normalisasi dengan font monospace/bold"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after  = Pt(2)
    pf.line_spacing = Pt(20)
    pf.left_indent  = Cm(0.5)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)
    run.font.bold = False
    return p

def add_table_entry(label, fields):
    """Tambah baris tabel normalisasi dengan label tabel dan fieldnya"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(1)
    pf.space_after  = Pt(1)
    pf.line_spacing = Pt(20)
    pf.left_indent  = Cm(0.5)
    # Label tabel (bold)
    run_label = p.add_run(label + ' = ')
    run_label.font.name = 'Times New Roman'
    run_label.font.size = Pt(11)
    run_label.font.bold = True
    # Fields
    run_fields = p.add_run(fields)
    run_fields.font.name = 'Times New Roman'
    run_fields.font.size = Pt(11)
    run_fields.font.bold = False
    return p

def add_separator():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    return p

# ============================================================
# JUDUL BAB
# ============================================================
p_title = add_heading('5.2.1.2  Perancangan Basis Data')
p_title.paragraph_format.space_before = Pt(0)

# ============================================================
# PEMBUKA
# ============================================================
add_paragraph(
    'Perancangan basis data pada sistem informasi Point of Sale (POS) dan Inventory '
    'DD Shoes Store dilakukan melalui proses normalisasi. Normalisasi adalah teknik '
    'formal dalam perancangan basis data relasional yang bertujuan untuk mengurangi '
    'redundansi data dan mencegah anomali data. Proses normalisasi dilakukan secara '
    'bertahap mulai dari Unnormalized Form (UNF) hingga Third Normal Form (3NF).',
    first_indent=1.25
)

# ============================================================
# PENJELASAN TEORI NORMALISASI (sesuai buku teks dari gambar)
# ============================================================
add_paragraph(
    'Berikut adalah penjelasan dari setiap tahapan normalisasi berdasarkan teori '
    'Database System Concepts (Silberschatz, Korth dan Sudarshan):',
    first_indent=1.25
)

# Teori UNF
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after  = Pt(4)
p.paragraph_format.line_spacing = Pt(24)
p.paragraph_format.left_indent  = Cm(1.25)
r1 = p.add_run('Unnormalized Form (UNF) : ')
set_font(r1, bold=True, size=12)
r2 = p.add_run(
    'Pada UNF, seluruh atribut akan dimasukkan dalam satu entitas. UNF juga tidak '
    'mewajibkan atribut yang dimasukkan harus sesuai dengan format tertentu. UNF '
    'memungkinkan adanya data berulang yang dapat menjadi masalah saat dilakukan '
    'manipulasi data. Hal ini biasa dikenal dengan anomali data.'
)
set_font(r2, bold=False, size=12)

# Teori 1NF
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after  = Pt(4)
p.paragraph_format.line_spacing = Pt(24)
p.paragraph_format.left_indent  = Cm(1.25)
r1 = p.add_run('First Normal Form (1NF) : ')
set_font(r1, bold=True, size=12)
r2 = p.add_run(
    '1NF mewajibkan setiap kolom memiliki nilai yang unik dan tidak ada kelompok '
    'data yang berulang dalam satu barisnya. Pada 1NF akan dilakukan pengelompokkan '
    'beberapa tipe data yang sejenis sehingga dapat mengatasi anomali data.'
)
set_font(r2, bold=False, size=12)

# Teori 2NF
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after  = Pt(4)
p.paragraph_format.line_spacing = Pt(24)
p.paragraph_format.left_indent  = Cm(1.25)
r1 = p.add_run('Second Normal Form (2NF) : ')
set_font(r1, bold=True, size=12)
r2 = p.add_run(
    '2NF mewajibkan setiap kolom dalam tabel hanya bergantung pada satu kunci primer '
    'dan tidak bergantung pada kolom lain dalam tabel. 2NF akan mewajibkan setiap '
    'tabel yang sudah dipisahkan memiliki kunci primer.'
)
set_font(r2, bold=False, size=12)

# Teori 3NF
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after  = Pt(4)
p.paragraph_format.line_spacing = Pt(24)
p.paragraph_format.left_indent  = Cm(1.25)
r1 = p.add_run('Third Normal Form (3NF) : ')
set_font(r1, bold=True, size=12)
r2 = p.add_run(
    '3NF mewajibkan setiap kolom dalam tabel hanya bergantung pada satu kunci primer '
    'dan tidak bergantung pada kolom lain yang bukan kunci dalam tabel. 3NF akan '
    'memisahkan atribut yang tidak bergantung langsung dengan kunci primer, tetapi '
    'bergantung pada atribut non key lainnya.'
)
set_font(r2, bold=False, size=12)

add_separator()

# ============================================================
# a. NORMALISASI
# ============================================================
add_heading('a.  Normalisasi')

add_paragraph(
    'Berikut ini adalah proses normalisasi basis data sistem POS DD Shoes Store '
    'yang dimulai dari bentuk tidak normal (UNF) hingga bentuk normal ketiga (3NF):',
    first_indent=1.25
)

# ============================================================
# UNF
# ============================================================
add_heading('1)  UNF (Unnormalized Form) / Bentuk Tidak Normal')

add_paragraph(
    'Pada tahap UNF, seluruh atribut dari semua entitas dalam sistem dikumpulkan '
    'menjadi satu kelompok data tanpa adanya pengelompokkan, kunci, maupun pemisahan '
    'antar entitas. Data masih bersifat mentah dan memungkinkan adanya atribut yang '
    'berulang (repeating group) serta anomali data.',
    first_indent=1.25
)

add_paragraph('Atribut UNF sistem DD Shoes POS:', bold=True, first_indent=1.25)

unf_text = (
    'username + password + is_active + last_login + date_joined + '
    'role + phone + '
    'category_name + brand_name + '
    'supplier_name + supplier_phone + supplier_notes + '
    'product_name + size + condition + description + buy_price + sell_price + '
    'stock + image + product_status + product_created_at + '
    'received_date + stockin_notes + stockin_created_at + '
    'stockin_qty + stockin_buy_price + '
    'cashier_username + subtotal_amount + discount_amount + total_amount + '
    'payment_method + cash_received + change_amount + transaction_date + '
    'transaction_status + void_reason + voided_at + '
    'trx_quantity + trx_sell_price + trx_subtotal + '
    'closing_date + system_cash_total + system_transfer_total + system_qris_total + '
    'actual_cash + cash_difference + closing_notes + is_locked + submitted_at + '
    'unlocked_at + unlock_reason + unlock_count + '
    'adj_quantity + adj_reason + adj_notes + adjusted_at'
)
add_formula_paragraph(unf_text)
add_separator()

# ============================================================
# 1NF
# ============================================================
add_heading('2)  1NF (First Normal Form) / Bentuk Normal Pertama')

add_paragraph(
    'Pada tahap 1NF, seluruh atribut dari UNF dikelompokkan berdasarkan kesamaan '
    'tipe data dan entitasnya masing-masing. Setiap tabel sudah memiliki atribut '
    'yang atomik (tidak ada kelompok berulang), namun belum memiliki kunci primer '
    'yang jelas dan masih terdapat ketergantungan parsial antar atribut.',
    first_indent=1.25
)

tables_1nf = [
    ('tb_user',
     'username + password + is_active + last_login + date_joined + role + phone'),
    ('tb_kategori',
     'category_name'),
    ('tb_merek',
     'brand_name'),
    ('tb_pemasok',
     'supplier_name + supplier_phone + supplier_notes'),
    ('tb_produk',
     'product_name + category_name + brand_name + size + condition + description + '
     'buy_price + sell_price + stock + image + product_status + product_created_at'),
    ('tb_barangmasuk',
     'supplier_name + received_by_username + received_date + stockin_notes + stockin_created_at + '
     'product_name + stockin_qty + stockin_buy_price'),
    ('tb_transaksi',
     'cashier_username + subtotal_amount + discount_amount + total_amount + '
     'payment_method + cash_received + change_amount + transaction_date + '
     'transaction_status + void_reason + voided_at + '
     'product_name + trx_quantity + trx_sell_price + trx_subtotal'),
    ('tb_closing',
     'cashier_username + closing_date + system_cash_total + system_transfer_total + '
     'system_qris_total + actual_cash + cash_difference + closing_notes + '
     'is_locked + submitted_at + unlocked_by_username + unlocked_at + '
     'unlock_reason + unlock_count'),
    ('tb_penyesuaianstok',
     'product_name + adjusted_by_username + adj_quantity + adj_reason + '
     'adj_notes + adjusted_at'),
]

for label, fields in tables_1nf:
    add_table_entry(label, fields)

add_separator()

# ============================================================
# 2NF
# ============================================================
add_heading('3)  2NF (Second Normal Form) / Bentuk Normal Kedua')

add_paragraph(
    'Pada tahap 2NF, tabel-tabel yang masih mengandung ketergantungan parsial '
    'dipisahkan menjadi tabel-tabel baru. Setiap tabel sudah memiliki kunci primer, '
    'dan setiap atribut non-kunci bergantung secara penuh pada kunci primer '
    'masing-masing tabelnya. Tabel yang memiliki atribut dari entitas lain '
    '(misalnya tb_produk yang masih menyimpan category_name dan brand_name) '
    'dipisah menjadi tabel tersendiri.',
    first_indent=1.25
)

tables_2nf = [
    ('tb_user',
     'user_id + username + password + is_active + last_login + date_joined'),
    ('tb_userprofile',
     'profile_id + user_id + role + phone'),
    ('tb_kategori',
     'category_id + category_name'),
    ('tb_merek',
     'brand_id + brand_name'),
    ('tb_pemasok',
     'supplier_id + supplier_name + supplier_phone + supplier_notes'),
    ('tb_produk',
     'product_id + category_id + brand_id + product_name + size + condition + '
     'description + buy_price + sell_price + stock + image + '
     'product_status + product_created_at'),
    ('tb_barangmasuk',
     'stockin_id + supplier_id + received_by + received_date + '
     'stockin_notes + stockin_created_at'),
    ('tb_detailbarangmasuk',
     'stockin_detail_id + stockin_id + product_id + stockin_qty + stockin_buy_price'),
    ('tb_transaksi',
     'transaction_id + cashier_id + subtotal_amount + discount_amount + '
     'total_amount + payment_method + cash_received + change_amount + '
     'transaction_date + transaction_status + voided_by + '
     'void_reason + voided_at'),
    ('tb_detailtransaksi',
     'trx_detail_id + transaction_id + product_id + trx_quantity + '
     'trx_sell_price + trx_subtotal'),
    ('tb_closing',
     'closing_id + cashier_id + closing_date + system_cash_total + '
     'system_transfer_total + system_qris_total + actual_cash + '
     'cash_difference + closing_notes + is_locked + submitted_at + '
     'unlocked_by + unlocked_at + unlock_reason + unlock_count'),
    ('tb_penyesuaianstok',
     'adjustment_id + product_id + adjusted_by + adj_quantity + '
     'adj_reason + adj_notes + adjusted_at'),
]

for label, fields in tables_2nf:
    add_table_entry(label, fields)

add_separator()

# ============================================================
# 3NF
# ============================================================
add_heading('4)  3NF (Third Normal Form) / Bentuk Normal Ketiga')

add_paragraph(
    'Pada tahap 3NF, seluruh ketergantungan transitif dihilangkan. Setiap atribut '
    'non-kunci hanya bergantung langsung pada kunci primer, bukan bergantung pada '
    'atribut non-kunci lainnya. Seluruh relasi antar tabel sudah dinyatakan '
    'secara eksplisit melalui Foreign Key (FK). Hasil 3NF ini merupakan struktur '
    'basis data final yang diimplementasikan pada sistem.',
    first_indent=1.25
)

tables_3nf = [
    ('tb_user',
     'user_id [PK] + username + password + is_active + last_login + date_joined'),
    ('tb_userprofile',
     'profile_id [PK] + user_id [FK] + role + phone'),
    ('tb_kategori',
     'category_id [PK] + category_name'),
    ('tb_merek',
     'brand_id [PK] + brand_name'),
    ('tb_pemasok',
     'supplier_id [PK] + supplier_name + supplier_phone + supplier_notes'),
    ('tb_produk',
     'product_id [PK] + category_id [FK] + brand_id [FK] + '
     'product_name + size + condition + description + '
     'buy_price + sell_price + stock + image + '
     'product_status + product_created_at'),
    ('tb_barangmasuk',
     'stockin_id [PK] + supplier_id [FK] + received_by [FK] + '
     'received_date + stockin_notes + stockin_created_at'),
    ('tb_detailbarangmasuk',
     'stockin_detail_id [PK] + stockin_id [FK] + product_id [FK] + '
     'stockin_qty + stockin_buy_price'),
    ('tb_transaksi',
     'transaction_id [PK] + cashier_id [FK] + '
     'subtotal_amount + discount_amount + total_amount + '
     'payment_method + cash_received + change_amount + '
     'transaction_date + transaction_status + '
     'voided_by [FK] + void_reason + voided_at'),
    ('tb_detailtransaksi',
     'trx_detail_id [PK] + transaction_id [FK] + product_id [FK] + '
     'trx_quantity + trx_sell_price + trx_subtotal'),
    ('tb_closing',
     'closing_id [PK] + cashier_id [FK] + closing_date + '
     'system_cash_total + system_transfer_total + system_qris_total + '
     'actual_cash + cash_difference + closing_notes + '
     'is_locked + submitted_at + '
     'unlocked_by [FK] + unlocked_at + unlock_reason + unlock_count'),
    ('tb_penyesuaianstok',
     'adjustment_id [PK] + product_id [FK] + adjusted_by [FK] + '
     'adj_quantity + adj_reason + adj_notes + adjusted_at'),
]

for label, fields in tables_3nf:
    add_table_entry(label, fields)

add_separator()

# ============================================================
# PENUTUP / KESIMPULAN NORMALISASI
# ============================================================
add_paragraph(
    'Berdasarkan proses normalisasi yang telah dilakukan dari tahap UNF hingga 3NF, '
    'diperoleh 12 (dua belas) tabel basis data yang saling berelasi. Setiap tabel '
    'telah memenuhi syarat bentuk normal ketiga (3NF) dimana setiap atribut non-kunci '
    'hanya bergantung secara langsung pada kunci primer masing-masing tabel. '
    'Struktur basis data ini selanjutnya diimplementasikan menggunakan Django ORM '
    'dengan database engine SQLite3.',
    first_indent=1.25
)

# ============================================================
# SIMPAN FILE
# ============================================================
output_path = r'c:\Users\IRVAN SUSANTO\pos_ddshoes2\pos_ddshoes\docs\Bab_5_2_1_2_Normalisasi_Final.docx'
doc.save(output_path)
print(f'File berhasil dibuat: {output_path}')
