import os

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
except ImportError:
    import sys
    print("Error: Pustaka python-docx belum terinstal.")
    print("Silakan jalankan perintah ini terlebih dahulu di terminal Anda:")
    print("pip install python-docx")
    sys.exit(1)

doc = Document()

# Judul
title = doc.add_heading('BAB 5\nHASIL PENELITIAN', level=1)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

doc.add_heading('5.1 Perancangan dan Pengembangan', level=2)
doc.add_heading('5.1.1 Fase Requirement Planning', level=3)

p1 = doc.add_paragraph(
    "Tahap perencanaan merupakan fase awal dalam pengembangan sistem yang mencakup serangkaian kegiatan identifikasi kebutuhan bersama pihak pengguna, guna memahami tujuan serta persyaratan yang diperlukan dalam membangun sistem informasi Point of Sale (POS) berbasis web pada UMKM DD Shoes Store Kubu Raya. Perencanaan dilaksanakan melalui pemahaman yang komprehensif terhadap kebutuhan operasional pengguna dalam menjalankan sistem, dengan tujuan mengidentifikasi secara menyeluruh aspek-aspek yang termasuk maupun yang berada di luar cakupan sistem yang akan dikembangkan, sehingga proses perancangan dapat berlangsung lebih terarah dan efisien dengan pendayagunaan sumber daya yang relevan secara optimal."
)
p1.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

p2 = doc.add_paragraph(
    "Data yang dimanfaatkan dalam tahap identifikasi kebutuhan diperoleh melalui serangkaian metode pengumpulan data, yang meliputi wawancara, observasi, serta dokumentasi yang dilaksanakan secara langsung di UMKM DD Shoes Store Kubu Raya. Dalam rangka menetapkan spesifikasi sistem yang akan dikembangkan secara tepat dan terukur, dilakukan perumusan terhadap dua kategorisasi kebutuhan sistem, yakni kebutuhan fungsional dan kebutuhan non-fungsional. Kebutuhan fungsional merujuk pada fungsi, proses, serta pengelolaan data yang harus dijalankan oleh sistem dalam menunjang kegiatan operasional, mencakup manajemen pengguna, pengelolaan data produk dan kategori, pemrosesan transaksi penjualan, serta penyajian laporan dan dashboard analitik. Sementara itu, kebutuhan non-fungsional berkaitan dengan kualitas dan karakteristik sistem dalam menjalankan fungsinya, yang meliputi aspek keamanan akses, kemudahan penggunaan antarmuka (usability), serta keandalan sistem dalam mendukung keberlangsungan aktivitas operasional UMKM DD Shoes Store Kubu Raya secara berkelanjutan."
)
p2.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

doc.add_heading('5.1.1.1 Kebutuhan Fungsional', level=4)
p3 = doc.add_paragraph(
    "Kebutuhan fungsional merupakan komponen yang memuat fungsi, aktivitas, serta pengelolaan data yang harus dijalankan oleh sistem agar dapat beroperasi secara optimal. Berdasarkan hasil perancangan sistem terbaru, adapun kebutuhan fungsional yang dirumuskan dalam pengembangan Sistem Informasi POS UMKM DD Shoes Store Kubu Raya adalah sebagai berikut:"
)
p3.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

doc.add_heading('a. Halaman Admin', level=5)
admin_features = [
    "Sistem wajib menyediakan fitur autentikasi login bagi administrator guna memperoleh akses terhadap keseluruhan halaman sistem.",
    "Sistem wajib menyediakan fitur logout bagi administrator untuk mengakhiri sesi akun yang sedang aktif.",
    "Sistem wajib menampilkan halaman dashboard yang memuat statistik jumlah kategori, jumlah produk, jumlah transaksi (Harian, Mingguan dan Bulanan), pendapatan (Harian, Mingguan dan Bulanan), grafik penjualan dalam format grafik garis, serta daftar lima produk dengan tingkat penjualan tertinggi.",
    "Sistem wajib menyediakan fitur bagi administrator untuk mengelola data merek (brands), yang meliputi penambahan, pengubahan, penampilan, dan penghapusan data merek.",
    "Sistem wajib menyediakan fitur bagi administrator untuk mengelola data kategori, yang meliputi penambahan, pengubahan, penampilan, dan penghapusan data kategori.",
    "Sistem wajib menyediakan fitur bagi administrator untuk mengelola data pemasok (suppliers), yang mencakup penambahan, pengubahan, penampilan, dan penghapusan data pemasok.",
    "Sistem wajib menyediakan fitur bagi administrator untuk mengelola data produk (katalog), yang mencakup penambahan melalui penerimaan barang masuk, pengubahan detail, penampilan, penghapusan data produk, serta penguncian status otomatis pada produk yang telah terjual.",
    "Sistem wajib menyediakan fitur bagi administrator untuk mencatat penerimaan barang masuk (stock in) beserta detail item produk dan referensi surat jalan dari pemasok.",
    "Sistem wajib menyediakan fitur bagi administrator untuk melakukan penyesuaian stok (stock adjustment) guna mencatat pengurangan stok akibat barang rusak atau hilang.",
    "Sistem wajib menyediakan fitur bagi administrator untuk mengunggah gambar pada saat pengelolaan data merek, kategori, dan produk.",
    "Sistem wajib menyediakan fitur bagi administrator untuk memproses transaksi penjualan melalui antarmuka POS yang interaktif.",
    "Sistem wajib menyediakan fitur bagi administrator untuk menetapkan diskon penjualan dalam bentuk persentase maupun nilai nominal pada saat berlangsungnya transaksi.",
    "Sistem wajib menyediakan fitur bagi administrator untuk menentukan metode pembayaran yang mencakup Tunai, Transfer BRI, dan QRIS.",
    "Sistem wajib melakukan kalkulasi uang kembalian secara otomatis pada transaksi dengan metode pembayaran Tunai.",
    "Sistem wajib menghasilkan kode struk (invoice) secara otomatis pada setiap transaksi penjualan yang diproses.",
    "Sistem wajib melakukan pengurangan stok produk secara otomatis pada saat transaksi penjualan berhasil dieksekusi.",
    "Sistem wajib menampilkan halaman laporan penjualan yang memuat riwayat seluruh transaksi beserta rincian item produk yang terjual.",
    "Sistem wajib menyediakan fitur bagi administrator untuk melakukan pencarian dan penyaringan data penjualan berdasarkan rentang periode waktu, nama kasir, metode pembayaran, serta merek produk.",
    "Sistem wajib menyediakan fitur bagi administrator untuk mencetak atau mengekspor laporan penjualan ke dalam format dokumen PDF.",
    "Sistem wajib menyediakan fitur bagi administrator untuk menampilkan pratinjau struk pembayaran (receipt) serta mencetaknya.",
    "Sistem wajib menyediakan fitur bagi administrator untuk membatalkan transaksi penjualan (retur) yang akan otomatis memulihkan stok produk dan menyesuaikan total pencatatan pendapatan secara tersinkronisasi.",
    "Sistem wajib menyediakan halaman laporan inventori guna memantau pergerakan stok, meliputi rincian total barang masuk, terjual, dan penyesuaian stok per produk.",
    "Sistem wajib menyediakan fitur bagi administrator untuk meninjau data pelaporan penutupan kasir (closing), serta memiliki wewenang untuk membuka ulang (unlock) akses form penutupan yang telah dikirimkan, dilengkapi dengan mekanisme pencatatan jejak audit (audit trail) secara sistematis.",
    "Sistem wajib menyediakan fitur bagi administrator untuk mengelola data akun pengguna (tambah, ubah, pengaturan status aktif, dan penugasan peran akses) melalui antarmuka manajemen pengguna yang terintegrasi di dalam sistem."
]
for i, feature in enumerate(admin_features, 1):
    doc.add_paragraph(f"{i}. {feature}", style='List Number')

doc.add_heading('b. Halaman Kasir', level=5)
kasir_features = [
    "Sistem wajib menyediakan fitur autentikasi login bagi kasir guna mengakses fungsionalitas sistem yang diizinkan sesuai dengan kewenangan perannya.",
    "Sistem wajib menyediakan fitur logout bagi kasir yang terintegrasi dengan algoritma peringatan kelalaian, yang secara proaktif mencegah proses keluar sesi apabila kasir belum menginput data pelaporan penutupan shift (closing) pada hari tersebut.",
    "Sistem wajib menampilkan halaman dashboard yang memuat ringkasan statistik dan grafik data penjualan yang dilakukan secara spesifik oleh akun kasir yang bersangkutan.",
    "Sistem wajib menyediakan halaman katalog produk dengan antarmuka khusus kasir guna meninjau daftar produk yang sedang aktif dan siap jual.",
    "Sistem wajib menyediakan fitur bagi kasir untuk memproses transaksi penjualan melalui antarmuka POS yang menampilkan interaksi pemilihan produk ke dalam keranjang.",
    "Sistem wajib menyediakan fitur bagi kasir untuk melakukan pencarian dan penyaringan produk berdasarkan nama, kategori, dan merek secara dinamis pada halaman POS.",
    "Sistem wajib menyediakan fitur bagi kasir untuk menetapkan besaran diskon penjualan pada saat berlangsungnya pemrosesan transaksi.",
    "Sistem wajib menyediakan fitur bagi kasir untuk menentukan jenis metode pembayaran, yaitu Tunai, Transfer BRI, atau QRIS.",
    "Sistem wajib menampilkan kalkulasi komputasi uang kembalian secara otomatis pada saat kasir memproses transaksi dengan metode pembayaran Tunai.",
    "Sistem wajib menyediakan fitur bagi kasir untuk menampilkan pratinjau antarmuka struk pembayaran serta fasilitas pencetakannya setelah transaksi berhasil diselesaikan.",
    "Sistem wajib menyediakan fitur formulir penutupan kasir (closing kasir) agar kasir dapat melaporkan nilai nominal serah terima uang tunai fisik yang ada di laci kasir pada setiap akhir shift operasional."
]
for i, feature in enumerate(kasir_features, 1):
    doc.add_paragraph(f"{i}. {feature}", style='List Number')

# Simpan ke folder dokumen
save_dir = r"c:\Users\IRVAN SUSANTO\pos_ddshoes2\pos_ddshoes\docs"
os.makedirs(save_dir, exist_ok=True)
file_path = os.path.join(save_dir, "Bab_5_1_1_Requirement_Planning_Revisi.docx")

doc.save(file_path)
print(f"Sukses! File skripsi berhasil dibuat di: {file_path}")
